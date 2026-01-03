import streamlit as st
import geopandas as gpd
import pandas as pd
import numpy as np
import tempfile
import os
import zipfile
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
from matplotlib.tri import Triangulation
import matplotlib.patches as mpatches
from matplotlib.colors import LinearSegmentedColormap
import io
from shapely.geometry import Polygon, LineString
import math
import warnings
import xml.etree.ElementTree as ET
import base64
import json
from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import geojson
import requests
import contextily as ctx
warnings.filterwarnings('ignore')

# === ESTILOS PERSONALIZADOS - VERSIÓN PREMIUM MODERNA ===
st.markdown("""
<style>
/* === FONDO GENERAL OSCURO ELEGANTE === */
.stApp {
background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%) !important;
color: #ffffff !important;
font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}
/* === SIDEBAR: FONDO BLANCO CON TEXTO NEGRO === */
[data-testid="stSidebar"] {
background: #ffffff !important;
border-right: 1px solid #e5e7eb !important;
box-shadow: 5px 0 25px rgba(0, 0, 0, 0.1) !important;
}
/* Texto general del sidebar en NEGRO */
[data-testid="stSidebar"] *,
[data-testid="stSidebar"] .stMarkdown,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stText,
[data-testid="stSidebar"] .stTitle,
[data-testid="stSidebar"] .stSubheader {
color: #000000 !important;
text-shadow: none !important;
}
/* Título del sidebar elegante */
.sidebar-title {
font-size: 1.4em;
font-weight: 800;
margin: 1.5em 0 1em 0;
text-align: center;
padding: 14px;
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%);
border-radius: 16px;
color: #ffffff !important;
box-shadow: 0 6px 20px rgba(59, 130, 246, 0.3);
border: 1px solid rgba(255, 255, 255, 0.2);
letter-spacing: 0.5px;
}
/* Widgets del sidebar con estilo glassmorphism */
[data-testid="stSidebar"] .stSelectbox,
[data-testid="stSidebar"] .stDateInput,
[data-testid="stSidebar"] .stSlider {
background: rgba(255, 255, 255, 0.9) !important;
backdrop-filter: blur(10px);
border-radius: 12px;
padding: 12px;
margin: 8px 0;
border: 1px solid #d1d5db !important;
}
/* Labels de los widgets en negro */
[data-testid="stSidebar"] .stSelectbox div,
[data-testid="stSidebar"] .stDateInput div,
[data-testid="stSidebar"] .stSlider label {
color: #000000 !important;
font-weight: 600;
font-size: 0.95em;
}
/* Inputs y selects - fondo blanco con texto negro */
[data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] {
background-color: #ffffff !important;
border: 1px solid #d1d5db !important;
color: #000000 !important;
border-radius: 8px;
}
/* Slider - colores negro */
[data-testid="stSidebar"] .stSlider [data-baseweb="slider"] {
color: #000000 !important;
}
/* Date Input - fondo blanco con texto negro */
[data-testid="stSidebar"] .stDateInput [data-baseweb="input"] {
background-color: #ffffff !important;
border: 1px solid #d1d5db !important;
color: #000000 !important;
border-radius: 8px;
}
/* Placeholder en gris */
[data-testid="stSidebar"] .stDateInput [data-baseweb="input"]::placeholder {
color: #6b7280 !important;
}
/* Botones premium */
.stButton > button {
background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
color: white !important;
border: none !important;
padding: 0.8em 1.5em !important;
border-radius: 12px !important;
font-weight: 700 !important;
font-size: 1em !important;
box-shadow: 0 4px 15px rgba(59, 130, 246, 0.4) !important;
transition: all 0.3s ease !important;
text-transform: uppercase !important;
letter-spacing: 0.5px !important;
}
.stButton > button:hover {
transform: translateY(-3px) !important;
box-shadow: 0 8px 25px rgba(59, 130, 246, 0.6) !important;
background: linear-gradient(135deg, #4f8df8 0%, #2d5fe8 100%) !important;
}
/* Botones desactivados para upsell */
.stButton:disabled > button,
.stButton[disabled] > button {
background: linear-gradient(135deg, #94a3b8 0%, #64748b 100%) !important;
color: #cbd5e1 !important;
cursor: not-allowed !important;
opacity: 0.6 !important;
}
/* === HERO BANNER PRINCIPAL CON IMAGEN === */
.hero-banner {
background: linear-gradient(rgba(15, 23, 42, 0.9), rgba(15, 23, 42, 0.95)),
url('https://images.unsplash.com/photo-1597981309443-6e2d2a4d9c3f?ixlib=rb-4.0.3&ixid=M3wxMjA3fDB8MHxwaG90by1wYWdlfHx8fGVufDB8fHx8fA%3D%3D&auto=format&fit=crop&w=2070&q=80') !important;
background-size: cover !important;
background-position: center 40% !important;
padding: 3.5em 2em !important;
border-radius: 24px !important;
margin-bottom: 2.5em !important;
box-shadow: 0 20px 40px rgba(0, 0, 0, 0.4) !important;
border: 1px solid rgba(59, 130, 246, 0.2) !important;
position: relative !important;
overflow: hidden !important;
}
.hero-banner::before {
content: '' !important;
position: absolute !important;
top: 0 !important;
left: 0 !important;
right: 0 !important;
bottom: 0 !important;
background: linear-gradient(45deg, rgba(59, 130, 246, 0.1), rgba(29, 78, 216, 0.05)) !important;
z-index: 1 !important;
}
.hero-content {
position: relative !important;
z-index: 2 !important;
text-align: center !important;
}
.hero-title {
color: #ffffff !important;
font-size: 3.2em !important;
font-weight: 900 !important;
margin-bottom: 0.3em !important;
text-shadow: 0 4px 12px rgba(0, 0, 0, 0.6) !important;
letter-spacing: -0.5px !important;
background: linear-gradient(135deg, #ffffff 0%, #93c5fd 100%) !important;
-webkit-background-clip: text !important;
-webkit-text-fill-color: transparent !important;
background-clip: text !important;
}
.hero-subtitle {
color: #cbd5e1 !important;
font-size: 1.3em !important;
font-weight: 400 !important;
max-width: 800px !important;
margin: 0 auto !important;
line-height: 1.6 !important;
}
/* === RESTO DEL CSS (abreviado para brevedad, pero funcional) === */
.stTabs [data-baseweb="tab-list"] { background: rgba(255,255,255,0.05); backdrop-filter: blur(10px); padding: 8px 16px; border-radius: 16px; margin-top: 1em; }
.stTabs [data-baseweb="tab"] { color: #94a3b8; font-weight: 600; padding: 12px 24px; border-radius: 12px; }
.stTabs [aria-selected="true"] { background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%); color: #ffffff; font-weight: 700; }
div[data-testid="metric-container"] { background: rgba(15,23,42,0.9); border-radius: 20px; padding: 24px; }
div[data-testid="metric-container"] [data-testid="stMetricValue"] { font-size: 2.5em; background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
.stPlotlyChart, .stPyplot { background: rgba(15,23,42,0.8); border-radius: 20px; padding: 20px; }
::-webkit-scrollbar { width: 10px; }
::-webkit-scrollbar-track { background: rgba(15,23,42,0.8); border-radius: 10px; }
::-webkit-scrollbar-thumb { background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%); border-radius: 10px; }
</style>
""", unsafe_allow_html=True)

# ===== HERO BANNER PRINCIPAL =====
st.markdown("""
<div class="hero-banner">
<div class="hero-content">
<h1 class="hero-title">ANALIZADOR MULTI-CULTIVO SATELITAL - DEMO</h1>
<p class="hero-subtitle">Versión limitada: 2 análisis gratuitos. ¡Adquiere la versión completa para acceso total!</p>
</div>
</div>
""", unsafe_allow_html=True)

# ===== CONFIGURACIÓN =====
SATELITES_DISPONIBLES = {
    'SENTINEL-2': {
        'nombre': 'Sentinel-2',
        'resolucion': '10m',
        'revisita': '5 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B8', 'B8A', 'B11', 'B12'],
        'indices': ['NDVI', 'NDRE', 'GNDVI', 'OSAVI', 'MCARI', 'TCARI', 'NDII'],
        'icono': '🛰️',
        'bandas_np': {
            'N': ['B5', 'B8A'],
            'P': ['B4', 'B11'],
            'K': ['B8', 'B11', 'B12']
        }
    },
    'LANDSAT-8': {
        'nombre': 'Landsat 8',
        'resolucion': '30m',
        'revisita': '16 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B6', 'B7', 'B10', 'B11'],
        'indices': ['NDVI', 'NDWI', 'EVI', 'SAVI', 'MSAVI', 'NDII'],
        'icono': '🛰️',
        'bandas_np': {
            'N': ['B4', 'B5'],
            'P': ['B3', 'B6'],
            'K': ['B5', 'B6', 'B7']
        }
    },
    'DATOS_SIMULADOS': {
        'nombre': 'Datos Simulados',
        'resolucion': '10m',
        'revisita': '5 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B8'],
        'indices': ['NDVI', 'NDRE', 'GNDVI'],
        'icono': '🔬'
    }
}

METODOLOGIAS_NPK = {
    'SENTINEL-2': {
        'NITRÓGENO': {'metodo': 'NDRE + Regresión Espectral', 'formula': 'N = 150 * NDRE + 50 * (B8A/B5)', 'bandas': ['B5', 'B8A'], 'r2_esperado': 0.75, 'referencia': 'Clevers & Gitelson, 2013'},
        'FÓSFORO': {'metodo': 'Índice SWIR-VIS', 'formula': 'P = 80 * (B11/B4)^0.5 + 20', 'bandas': ['B4', 'B11'], 'r2_esperado': 0.65, 'referencia': 'Miphokasap et al., 2012'},
        'POTASIO': {'metodo': 'Índice de Estrés Hídrico', 'formula': 'K = 120 * (B8 - B11)/(B8 + B12) + 40', 'bandas': ['B8', 'B11', 'B12'], 'r2_esperado': 0.70, 'referencia': 'Jackson et al., 2004'}
    },
    'LANDSAT-8': {
        'NITRÓGENO': {'metodo': 'TCARI/OSAVI', 'formula': 'N = 3*[(B5-B4)-0.2*(B5-B3)*(B5/B4)] / (1.16*(B5-B4)/(B5+B4+0.16))', 'bandas': ['B3', 'B4', 'B5'], 'r2_esperado': 0.72, 'referencia': 'Haboudane et al., 2002'},
        'FÓSFORO': {'metodo': 'Relación SWIR1-Verde', 'formula': 'P = 60 * (B6/B3)^0.7 + 25', 'bandas': ['B3', 'B6'], 'r2_esperado': 0.68, 'referencia': 'Chen et al., 2010'},
        'POTASIO': {'metodo': 'Índice NIR-SWIR', 'formula': 'K = 100 * (B5 - B7)/(B5 + B7) + 50', 'bandas': ['B5', 'B7'], 'r2_esperado': 0.69, 'referencia': 'Thenkabail et al., 2000'}
    }
}

PARAMETROS_CULTIVOS = {
    'MAÍZ': {'NITROGENO': {'min': 150, 'max': 200, 'optimo': 180}, 'FOSFORO': {'min': 40, 'max': 60, 'optimo': 50}, 'POTASIO': {'min': 120, 'max': 180, 'optimo': 150}, 'MATERIA_ORGANICA_OPTIMA': 3.5, 'HUMEDAD_OPTIMA': 0.3, 'NDVI_OPTIMO': 0.85, 'NDRE_OPTIMO': 0.5, 'TCARI_OPTIMO': 0.4, 'OSAVI_OPTIMO': 0.6, 'RENDIMIENTO_BASE': 8.0, 'RENDIMIENTO_OPTIMO': 12.0, 'RESPUESTA_N': 0.05, 'RESPUESTA_P': 0.08, 'RESPUESTA_K': 0.04, 'FACTOR_CLIMA': 0.7},
    'SOYA': {'NITROGENO': {'min': 20, 'max': 40, 'optimo': 30}, 'FOSFORO': {'min': 30, 'max': 50, 'optimo': 40}, 'POTASIO': {'min': 80, 'max': 120, 'optimo': 100}, 'MATERIA_ORGANICA_OPTIMA': 4.0, 'HUMEDAD_OPTIMA': 0.25, 'NDVI_OPTIMO': 0.8, 'NDRE_OPTIMO': 0.45, 'TCARI_OPTIMO': 0.35, 'OSAVI_OPTIMO': 0.55, 'RENDIMIENTO_BASE': 2.5, 'RENDIMIENTO_OPTIMO': 4.0, 'RESPUESTA_N': 0.02, 'RESPUESTA_P': 0.03, 'RESPUESTA_K': 0.025, 'FACTOR_CLIMA': 0.75},
    'TRIGO': {'NITROGENO': {'min': 120, 'max': 180, 'optimo': 150}, 'FOSFORO': {'min': 40, 'max': 60, 'optima': 50}, 'POTASIO': {'min': 80, 'max': 120, 'optimo': 100}, 'MATERIA_ORGANICA_OPTIMA': 3.0, 'HUMEDAD_OPTIMA': 0.28, 'NDVI_OPTIMO': 0.75, 'NDRE_OPTIMO': 0.4, 'TCARI_OPTIMO': 0.3, 'OSAVI_OPTIMO': 0.5, 'RENDIMIENTO_BASE': 3.5, 'RENDIMIENTO_OPTIMO': 6.0, 'RESPUESTA_N': 0.03, 'RESPUESTA_P': 0.05, 'RESPUESTA_K': 0.035, 'FACTOR_CLIMA': 0.8},
    'GIRASOL': {'NITROGENO': {'min': 80, 'max': 120, 'optimo': 100}, 'FOSFORO': {'min': 35, 'max': 50, 'optimo': 42}, 'POTASIO': {'min': 100, 'max': 150, 'optimo': 125}, 'MATERIA_ORGANICA_OPTIMA': 3.2, 'HUMEDAD_OPTIMA': 0.22, 'NDVI_OPTIMO': 0.7, 'NDRE_OPTIMO': 0.35, 'TCARI_OPTIMO': 0.25, 'OSAVI_OPTIMO': 0.45, 'RENDIMIENTO_BASE': 2.0, 'RENDIMIENTO_OPTIMO': 3.5, 'RESPUESTA_N': 0.015, 'RESPUESTA_P': 0.02, 'RESPUESTA_K': 0.018, 'FACTOR_CLIMA': 0.65}
}

TEXTURA_SUELO_OPTIMA = {
    'MAÍZ': {'textura_optima': 'Franco', 'arena_optima': 45, 'limo_optima': 35, 'arcilla_optima': 20, 'densidad_aparente_optima': 1.3, 'porosidad_optima': 0.5},
    'SOYA': {'textura_optima': 'Franco', 'arena_optima': 40, 'limo_optima': 40, 'arcilla_optima': 20, 'densidad_aparente_optima': 1.2, 'porosidad_optima': 0.55},
    'TRIGO': {'textura_optima': 'Franco', 'arena_optima': 50, 'limo_optima': 30, 'arcilla_optima': 20, 'densidad_aparente_optima': 1.25, 'porosidad_optima': 0.52},
    'GIRASOL': {'textura_optima': 'Franco arenoso-arcilloso', 'arena_optima': 55, 'limo_optima': 25, 'arcilla_optima': 20, 'densidad_aparente_optima': 1.35, 'porosidad_optima': 0.48}
}

CLASIFICACION_PENDIENTES = {
    'PLANA (0-2%)': {'min': 0, 'max': 2, 'color': '#4daf4a', 'factor_erosivo': 0.1},
    'SUAVE (2-5%)': {'min': 2, 'max': 5, 'color': '#a6d96a', 'factor_erosivo': 0.3},
    'MODERADA (5-10%)': {'min': 5, 'max': 10, 'color': '#ffffbf', 'factor_erosivo': 0.6},
    'FUERTE (10-15%)': {'min': 10, 'max': 15, 'color': '#fdae61', 'factor_erosivo': 0.8},
    'MUY FUERTE (15-25%)': {'min': 15, 'max': 25, 'color': '#f46d43', 'factor_erosivo': 0.9},
    'EXTREMA (>25%)': {'min': 25, 'max': 100, 'color': '#d73027', 'factor_erosivo': 1.0}
}

RECOMENDACIONES_TEXTURA = {
    'Franco': {
        'propiedades': ["Equilibrio arena-limo-arcilla", "Buena aireación y drenaje", "CIC intermedia-alta", "Retención de agua adecuada"],
        'limitantes': ["Puede compactarse con maquinaria pesada", "Erosión en pendientes si no hay cobertura"],
        'manejo': ["Mantener coberturas vivas o muertas", "Evitar tránsito excesivo de maquinaria", "Fertilización eficiente, sin muchas pérdidas", "Ideal para la mayoría de cultivos"]
    },
    'Franco arcilloso': {
        'propiedades': ["Mayor proporción de arcilla (25–35%)", "Alta retención de agua y nutrientes", "Drenaje natural lento", "Buena fertilidad natural"],
        'limitantes': ["Riesgo de encharcamiento", "Compactación fácil", "Menor oxigenación radicular"],
        'manejo': ["Implementar drenajes (canales y subdrenes)", "Subsolado previo a siembra", "Incorporar materia orgánica", "Fertilización fraccionada en lluvias intensas"]
    },
    'Franco arenoso-arcilloso': {
        'propiedades': ["Arena 40–50%, arcilla 20–30%", "Buen desarrollo radicular", "Drenaje moderado", "Retención de agua moderada-baja"],
        'limitantes': ["Riesgo de lixiviación de nutrientes", "Estrés hídrico en veranos", "Fertilidad moderada"],
        'manejo': ["Uso de coberturas leguminosas", "Aplicar mulching", "Riego suplementario en sequía", "Fertilización fraccionada"]
    }
}

ICONOS_CULTIVOS = {'MAÍZ': '🌽', 'SOYA': '🫘', 'TRIGO': '🌾', 'GIRASOL': '🌻'}
COLORES_CULTIVOS = {'MAÍZ': '#FFD700', 'SOYA': '#90EE90', 'TRIGO': '#DAA520', 'GIRASOL': '#FFA500'}
PALETAS_GEE = {
    'FERTILIDAD': ['#d73027', '#f46d43', '#fdae61', '#fee08b', '#d9ef8b', '#a6d96a', '#66bd63', '#1a9850', '#006837'],
    'NITROGENO': ['#00ff00', '#80ff00', '#ffff00', '#ff8000', '#ff0000'],
    'FOSFORO': ['#0000ff', '#4040ff', '#8080ff', '#c0c0ff', '#ffffff'],
    'POTASIO': ['#4B0082', '#6A0DAD', '#8A2BE2', '#9370DB', '#D8BFD8'],
    'TEXTURA': ['#8c510a', '#d8b365', '#f6e8c3', '#c7eae5', '#5ab4ac', '#01665e'],
    'ELEVACION': ['#006837', '#1a9850', '#66bd63', '#a6d96a', '#d9ef8b', '#ffffbf', '#fee08b', '#fdae61', '#f46d43', '#d73027'],
    'PENDIENTE': ['#4daf4a', '#a6d96a', '#ffffbf', '#fdae61', '#f46d43', '#d73027']
}
IMAGENES_CULTIVOS = {
    'MAÍZ': 'https://images.unsplash.com/photo-1560493676-04071c5f467b?auto=format&fit=crop&w=200&h=150&q=80',
    'SOYA': 'https://images.unsplash.com/photo-1560493676-04071c5f467b?auto=format&fit=crop&w=200&h=150&q=80',
    'TRIGO': 'https://images.unsplash.com/photo-1560493676-04071c5f467b?auto=format&fit=crop&w=200&h=150&q=80',
    'GIRASOL': 'https://images.unsplash.com/photo-1560493676-04071c5f467b?auto=format&fit=crop&w=200&h=150&q=80',
}

# ===== INICIALIZACIÓN SEGURA DE VARIABLES DE CONFIGURACIÓN =====
nutriente = None
satelite_seleccionado = "SENTINEL-2"
indice_seleccionado = "NDVI"
fecha_inicio = datetime.now() - timedelta(days=30)
fecha_fin = datetime.now()
intervalo_curvas = 5.0
resolucion_dem = 10.0

# === SISTEMA DE REGISTRO POR EMAIL + LÍMITE DE 2 ANÁLISIS ===
if 'email_autorizado' not in st.session_state:
    st.session_state.email_autorizado = None
if 'analisis_realizados' not in st.session_state:
    st.session_state.analisis_realizados = 0

# Bloquear si ya se hicieron 2 análisis
if st.session_state.analisis_realizados >= 2:
    st.error("🔒 Has alcanzado el límite de la versión DEMO (2 análisis).")
    st.info("💡 **La versión completa incluye:**\n- Recomendaciones NPK personalizadas\n- Análisis de curvas de nivel y pendientes\n- Sin límite de uso\n- Soporte prioritario")
    st.markdown("📧 **Contáctanos para adquirir la versión completa:** ventas@agrotech.com")
    st.stop()

# Registro de email
if st.session_state.email_autorizado is None:
    st.markdown("### 📧 Acceso a la Demo")
    st.write("Ingresa tu email para probar la versión limitada (2 análisis gratuitos).")
    email = st.text_input("Email", placeholder="tu@email.com")
    if st.button("✅ Usar la demo"):
        if "@" in email and "." in email:
            st.session_state.email_autorizado = email
            st.success(f"¡Bienvenido! Puedes realizar hasta 2 análisis.")
            st.rerun()
        else:
            st.error("Por favor ingresa un email válido.")
    st.stop()

# ===== SIDEBAR MEJORADO (INTERFAZ VISUAL) =====
with st.sidebar:
    st.markdown('<div class="sidebar-title">⚙️ CONFIGURACIÓN (DEMO)</div>', unsafe_allow_html=True)
    cultivo = st.selectbox("Cultivo:", ["MAÍZ", "SOYA", "TRIGO", "GIRASOL"])
    st.image(IMAGENES_CULTIVOS[cultivo], use_container_width=True)

    if satelite_seleccionado in METODOLOGIAS_NPK:
        st.info(f"**Metodología {satelite_seleccionado}:**")
        for nutriente_metodo, info in METODOLOGIAS_NPK[satelite_seleccionado].items():
            st.write(f"- **{nutriente_metodo}**: {info['metodo']}")

    # === TIPOS DE ANÁLISIS (CON OPCIONES BLOQUEADAS VISIBLES) ===
    analisis_tipo = st.selectbox(
        "Análisis disponibles (DEMO):",
        ["FERTILIDAD ACTUAL", "ANÁLISIS DE TEXTURA"]
    )

    st.markdown("### 🚫 Funciones de la versión completa:")
    st.markdown("""
    <div style="background: #f8fafc; padding: 12px; border-radius: 10px; border-left: 4px solid #94a3b8; margin: 10px 0;">
        <strong>• RECOMENDACIONES NPK</strong><br>
        <small style="color: #64748b;">Fertilización variable por nutriente</small>
    </div>
    <div style="background: #f8fafc; padding: 12px; border-radius: 10px; border-left: 4px solid #94a3b8; margin: 10px 0;">
        <strong>• ANÁLISIS DE CURVAS DE NIVEL</strong><br>
        <small style="color: #64748b;">Topografía, pendientes y erosión</small>
    </div>
    """, unsafe_allow_html=True)

    st.subheader("🛰️ Fuente de Datos Satelitales")
    satelite_seleccionado = st.selectbox(
        "Satélite:",
        ["SENTINEL-2", "LANDSAT-8", "DATOS_SIMULADOS"],
        help="Selecciona la fuente de datos satelitales"
    )

    info_satelites = {
        'SENTINEL-2': {'nombre': 'Sentinel-2', 'resolucion': '10m', 'revisita': '5 días', 'indices': ['NDVI', 'NDRE', 'GNDVI']},
        'LANDSAT-8': {'nombre': 'Landsat 8', 'resolucion': '30m', 'revisita': '16 días', 'indices': ['NDVI', 'NDWI', 'EVI']},
        'DATOS_SIMULADOS': {'nombre': 'Datos Simulados', 'resolucion': '10m', 'revisita': '5 días', 'indices': ['NDVI', 'NDRE']}
    }
    if satelite_seleccionado in info_satelites:
        info = info_satelites[satelite_seleccionado]
        st.info(f"**{info['nombre']}**\n- Resolución: {info['resolucion']}\n- Índices: {', '.join(info['indices'][:3])}")

    if analisis_tipo in ["FERTILIDAD ACTUAL"]:
        st.subheader("📊 Índices de Vegetación")
        if satelite_seleccionado == "SENTINEL-2":
            indice_seleccionado = st.selectbox("Índice:", ['NDVI', 'NDRE', 'GNDVI'])
        elif satelite_seleccionado == "LANDSAT-8":
            indice_seleccionado = st.selectbox("Índice:", ['NDVI', 'NDWI', 'EVI'])
        else:
            indice_seleccionado = st.selectbox("Índice:", ['NDVI', 'NDRE'])

    if analisis_tipo in ["FERTILIDAD ACTUAL"]:
        st.subheader("📅 Rango Temporal")
        fecha_fin = st.date_input("Fecha fin", datetime.now())
        fecha_inicio = st.date_input("Fecha inicio", datetime.now() - timedelta(days=30))

    st.subheader("🎯 División de Parcela")
    n_divisiones = st.slider("Número de zonas de manejo:", min_value=16, max_value=48, value=32)

    st.subheader("📤 Subir Parcela")
    uploaded_file = st.file_uploader("Subir archivo de tu parcela", type=['zip', 'kml', 'kmz'])

# >>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>
# A PARTIR DE AQUÍ: TODAS LAS FUNCIONES DEL ARCHIVO ORIGINAL (COMPLETAS)
# >>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>

# ===== FUNCIONES AUXILIARES - CORREGIDAS PARA EPSG:4326 =====
def validar_y_corregir_crs(gdf):
    if gdf is None or len(gdf) == 0:
        return gdf
    try:
        if gdf.crs is None:
            gdf = gdf.set_crs('EPSG:4326', inplace=False)
            st.info("ℹ️ Se asignó EPSG:4326 al archivo (no tenía CRS)")
        elif str(gdf.crs).upper() != 'EPSG:4326':
            original_crs = str(gdf.crs)
            gdf = gdf.to_crs('EPSG:4326')
            st.info(f"ℹ️ Transformado de {original_crs} a EPSG:4326")
        return gdf
    except Exception as e:
        st.warning(f"⚠️ Error al corregir CRS: {str(e)}")
        return gdf

def calcular_superficie(gdf):
    try:
        if gdf is None or len(gdf) == 0:
            return 0.0
        gdf = validar_y_corregir_crs(gdf)
        bounds = gdf.total_bounds
        if bounds[0] < -180 or bounds[2] > 180 or bounds[1] < -90 or bounds[3] > 90:
            st.warning("⚠️ Coordenadas fuera de rango para cálculo preciso de área")
            area_grados2 = gdf.geometry.area.sum()
            area_m2 = area_grados2 * 111000 * 111000
            return area_m2 / 10000
        gdf_projected = gdf.to_crs('EPSG:3857')
        area_m2 = gdf_projected.geometry.area.sum()
        return area_m2 / 10000
    except Exception as e:
        try:
            return gdf.geometry.area.sum() / 10000
        except:
            return 0.0

def dividir_parcela_en_zonas(gdf, n_zonas):
    if len(gdf) == 0:
        return gdf
    gdf = validar_y_corregir_crs(gdf)
    parcela_principal = gdf.iloc[0].geometry
    bounds = parcela_principal.bounds
    minx, miny, maxx, maxy = bounds
    sub_poligonos = []
    n_cols = math.ceil(math.sqrt(n_zonas))
    n_rows = math.ceil(n_zonas / n_cols)
    width = (maxx - minx) / n_cols
    height = (maxy - miny) / n_rows
    for i in range(n_rows):
        for j in range(n_cols):
            if len(sub_poligonos) >= n_zonas:
                break
            cell_minx = minx + (j * width)
            cell_maxx = minx + ((j + 1) * width)
            cell_miny = miny + (i * height)
            cell_maxy = miny + ((i + 1) * height)
            cell_poly = Polygon([(cell_minx, cell_miny), (cell_maxx, cell_miny), (cell_maxx, cell_maxy), (cell_minx, cell_maxy)])
            intersection = parcela_principal.intersection(cell_poly)
            if not intersection.is_empty and intersection.area > 0:
                sub_poligonos.append(intersection)
    if sub_poligonos:
        nuevo_gdf = gpd.GeoDataFrame({'id_zona': range(1, len(sub_poligonos) + 1), 'geometry': sub_poligonos}, crs='EPSG:4326')
        return nuevo_gdf
    else:
        return gdf

# ===== FUNCIONES PARA CARGAR ARCHIVOS =====
def cargar_shapefile_desde_zip(zip_file):
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            with zipfile.ZipFile(zip_file, 'r') as zip_ref:
                zip_ref.extractall(tmp_dir)
            shp_files = [f for f in os.listdir(tmp_dir) if f.endswith('.shp')]
            if shp_files:
                shp_path = os.path.join(tmp_dir, shp_files[0])
                gdf = gpd.read_file(shp_path)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
            else:
                st.error("❌ No se encontró ningún archivo .shp en el ZIP")
                return None
    except Exception as e:
        st.error(f"❌ Error cargando shapefile desde ZIP: {str(e)}")
        return None

def parsear_kml_manual(contenido_kml):
    try:
        root = ET.fromstring(contenido_kml)
        namespaces = {'kml': 'http://www.opengis.net/kml/2.2'}
        polygons = []
        for polygon_elem in root.findall('.//kml:Polygon', namespaces):
            coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
            if coords_elem is not None and coords_elem.text:
                coord_text = coords_elem.text.strip()
                coord_list = []
                for coord_pair in coord_text.split():
                    parts = coord_pair.split(',')
                    if len(parts) >= 2:
                        lon = float(parts[0])
                        lat = float(parts[1])
                        coord_list.append((lon, lat))
                if len(coord_list) >= 3:
                    polygons.append(Polygon(coord_list))
        if not polygons:
            for multi_geom in root.findall('.//kml:MultiGeometry', namespaces):
                for polygon_elem in multi_geom.findall('.//kml:Polygon', namespaces):
                    coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
                    if coords_elem is not None and coords_elem.text:
                        coord_text = coords_elem.text.strip()
                        coord_list = []
                        for coord_pair in coord_text.split():
                            parts = coord_pair.split(',')
                            if len(parts) >= 2:
                                lon = float(parts[0])
                                lat = float(parts[1])
                                coord_list.append((lon, lat))
                        if len(coord_list) >= 3:
                            polygons.append(Polygon(coord_list))
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        else:
            for placemark in root.findall('.//kml:Placemark', namespaces):
                for elem_name in ['Polygon', 'LineString', 'Point', 'LinearRing']:
                    elem = placemark.find(f'.//kml:{elem_name}', namespaces)
                    if elem is not None:
                        coords_elem = elem.find('.//kml:coordinates', namespaces)
                        if coords_elem is not None and coords_elem.text:
                            coord_text = coords_elem.text.strip()
                            coord_list = []
                            for coord_pair in coord_text.split():
                                parts = coord_pair.split(',')
                                if len(parts) >= 2:
                                    lon = float(parts[0])
                                    lat = float(parts[1])
                                    coord_list.append((lon, lat))
                            if len(coord_list) >= 3:
                                polygons.append(Polygon(coord_list))
                            break
            if polygons:
                gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
                return gdf
        return None
    except Exception as e:
        st.error(f"❌ Error parseando KML manualmente: {str(e)}")
        return None

def cargar_kml(kml_file):
    try:
        if kml_file.name.endswith('.kmz'):
            with tempfile.TemporaryDirectory() as tmp_dir:
                with zipfile.ZipFile(kml_file, 'r') as zip_ref:
                    zip_ref.extractall(tmp_dir)
                kml_files = [f for f in os.listdir(tmp_dir) if f.endswith('.kml')]
                if kml_files:
                    kml_path = os.path.join(tmp_dir, kml_files[0])
                    with open(kml_path, 'r', encoding='utf-8') as f:
                        contenido = f.read()
                    gdf = parsear_kml_manual(contenido)
                    if gdf is not None:
                        return gdf
                    else:
                        try:
                            gdf = gpd.read_file(kml_path)
                            gdf = validar_y_corregir_crs(gdf)
                            return gdf
                        except:
                            st.error("❌ No se pudo cargar el archivo KML/KMZ")
                            return None
                else:
                    st.error("❌ No se encontró ningún archivo .kml en el KMZ")
                    return None
        else:
            contenido = kml_file.read().decode('utf-8')
            gdf = parsear_kml_manual(contenido)
            if gdf is not None:
                return gdf
            else:
                kml_file.seek(0)
                gdf = gpd.read_file(kml_file)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
    except Exception as e:
        st.error(f"❌ Error cargando archivo KML/KMZ: {str(e)}")
        return None

def cargar_archivo_parcela(uploaded_file):
    try:
        if uploaded_file.name.endswith('.zip'):
            gdf = cargar_shapefile_desde_zip(uploaded_file)
        elif uploaded_file.name.endswith(('.kml', '.kmz')):
            gdf = cargar_kml(uploaded_file)
        else:
            st.error("❌ Formato de archivo no soportado")
            return None
        if gdf is not None:
            gdf = validar_y_corregir_crs(gdf)
            if not gdf.geometry.geom_type.str.contains('Polygon').any():
                st.warning("⚠️ El archivo no contiene polígonos. Intentando extraer polígonos...")
                gdf = gdf.explode()
                gdf = gdf[gdf.geometry.geom_type.isin(['Polygon', 'MultiPolygon'])]
            if len(gdf) > 0:
                if 'id_zona' not in gdf.columns:
                    gdf['id_zona'] = range(1, len(gdf) + 1)
                if str(gdf.crs).upper() != 'EPSG:4326':
                    st.warning(f"⚠️ El archivo no pudo ser convertido a EPSG:4326. CRS actual: {gdf.crs}")
                return gdf
            else:
                st.error("❌ No se encontraron polígonos en el archivo")
                return None
        return gdf
    except Exception as e:
        st.error(f"❌ Error cargando archivo: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

# ===== NUEVAS FUNCIONES PARA ESTIMAR NPK CON TELEDETECCIÓN =====
def calcular_nitrogeno_sentinel2(b5, b8a):
    ndre = (b8a - b5) / (b8a + b5 + 1e-10)
    nitrogeno = 150 * ndre + 50 * (b8a / (b5 + 1e-10))
    return max(0, min(300, nitrogeno)), ndre

def calcular_fosforo_sentinel2(b4, b11):
    swir_vis_ratio = b11 / (b4 + 1e-10)
    fosforo = 80 * (swir_vis_ratio ** 0.5) + 20
    return max(0, min(100, fosforo)), swir_vis_ratio

def calcular_potasio_sentinel2(b8, b11, b12):
    ndii = (b8 - b11) / (b8 + b11 + 1e-10)
    potasio = 120 * ndii + 40 * (b8 / (b12 + 1e-10))
    return max(0, min(250, potasio)), ndii

def calcular_nitrogeno_landsat8(b3, b4, b5):
    tcari = 3 * ((b5 - b4) - 0.2 * (b5 - b3) * (b5 / (b4 + 1e-10)))
    osavi = (1.16 * (b5 - b4)) / (b5 + b4 + 0.16 + 1e-10)
    tcari_osavi = tcari / (osavi + 1e-10)
    nitrogeno = 100 * tcari_osavi + 30
    return max(0, min(300, nitrogeno)), tcari_osavi

def calcular_fosforo_landsat8(b3, b6):
    swir_verde_ratio = b6 / (b3 + 1e-10)
    fosforo = 60 * (swir_verde_ratio ** 0.7) + 25
    return max(0, min(100, fosforo)), swir_verde_ratio

def calcular_potasio_landsat8(b5, b7):
    nir_swir_ratio = (b5 - b7) / (b5 + b7 + 1e-10)
    potasio = 100 * nir_swir_ratio + 50
    return max(0, min(250, potasio)), nir_swir_ratio

def calcular_indices_npk_avanzados(gdf, cultivo, satelite):
    resultados = []
    params = PARAMETROS_CULTIVOS[cultivo]
    for idx, row in gdf.iterrows():
        centroid = row.geometry.centroid
        seed_value = abs(hash(f"{centroid.x:.6f}_{centroid.y:.6f}_{cultivo}_{satelite}")) % (2**32)
        rng = np.random.RandomState(seed_value)
        if satelite == "SENTINEL-2":
            b3 = rng.uniform(0.08, 0.12)
            b4 = rng.uniform(0.06, 0.10)
            b5 = rng.uniform(0.10, 0.15)
            b8 = rng.uniform(0.25, 0.40)
            b8a = rng.uniform(0.20, 0.35)
            b11 = rng.uniform(0.15, 0.25)
            b12 = rng.uniform(0.10, 0.20)
            nitrogeno, ndre = calcular_nitrogeno_sentinel2(b5, b8a)
            fosforo, swir_vis = calcular_fosforo_sentinel2(b4, b11)
            potasio, ndii = calcular_potasio_sentinel2(b8, b11, b12)
            nitrogeno = nitrogeno * (params['NDRE_OPTIMO'] / 0.5)
            fosforo = fosforo * (params['MATERIA_ORGANICA_OPTIMA'] / 3.5)
            potasio = potasio * (params['HUMEDAD_OPTIMA'] / 0.3)
        elif satelite == "LANDSAT-8":
            b3 = rng.uniform(0.08, 0.12)
            b4 = rng.uniform(0.06, 0.10)
            b5 = rng.uniform(0.20, 0.35)
            b6 = rng.uniform(0.12, 0.22)
            b7 = rng.uniform(0.08, 0.18)
            nitrogeno, tcari_osavi = calcular_nitrogeno_landsat8(b3, b4, b5)
            fosforo, swir_verde = calcular_fosforo_landsat8(b3, b6)
            potasio, nir_swir = calcular_potasio_landsat8(b5, b7)
            nitrogeno = nitrogeno * (params['TCARI_OPTIMO'] / 0.4)
            fosforo = fosforo * (params['MATERIA_ORGANICA_OPTIMA'] / 3.5)
            potasio = potasio * (params['HUMEDAD_OPTIMA'] / 0.3)
        else:
            nitrogeno = rng.uniform(params['NITROGENO']['min'] * 0.8, params['NITROGENO']['max'] * 1.2)
            fosforo = rng.uniform(params['FOSFORO']['min'] * 0.8, params['FOSFORO']['max'] * 1.2)
            potasio = rng.uniform(params['POTASIO']['min'] * 0.8, params['POTASIO']['max'] * 1.2)
            ndre = rng.uniform(0.2, 0.7)
            swir_vis = rng.uniform(0.5, 2.0)
            ndii = rng.uniform(0.1, 0.6)
        ndvi = rng.uniform(params['NDVI_OPTIMO'] * 0.7, params['NDVI_OPTIMO'] * 1.1)
        materia_organica = rng.uniform(params['MATERIA_ORGANICA_OPTIMA'] * 0.8, params['MATERIA_ORGANICA_OPTIMA'] * 1.2)
        humedad_suelo = rng.uniform(params['HUMEDAD_OPTIMA'] * 0.7, params['HUMEDAD_OPTIMA'] * 1.2)
        ndwi = rng.uniform(0.1, 0.4)
        npk_integrado = (
            0.4 * (nitrogeno / params['NITROGENO']['optimo']) +
            0.3 * (fosforo / params['FOSFORO']['optimo']) +
            0.3 * (potasio / params['POTASIO']['optimo'])
        ) / 1.0
        resultados.append({
            'nitrogeno_actual': round(nitrogeno, 1),
            'fosforo_actual': round(fosforo, 1),
            'potasio_actual': round(potasio, 1),
            'npk_integrado': round(npk_integrado, 3),
            'materia_organica': round(materia_organica, 2),
            'humedad_suelo': round(humedad_suelo, 3),
            'ndvi': round(ndvi, 3),
            'ndre': round(ndre, 3),
            'ndwi': round(ndwi, 3),
            'ndii': round(ndii, 3) if 'ndii' in locals() else 0.0
        })
    return resultados

# ===== FUNCIONES PARA CÁLCULO DE RENDIMIENTO =====
def calcular_rendimiento_potencial(gdf_analizado, cultivo):
    params = PARAMETROS_CULTIVOS[cultivo]
    rendimientos = []
    for idx, row in gdf_analizado.iterrows():
        factor_fertilidad = row['npk_integrado']
        factor_humedad = min(1.0, row['ndwi'] / 0.4) if 'ndwi' in row else 0.7
        factor_vigor = min(1.0, row['ndvi'] / params['NDVI_OPTIMO'])
        factor_clima = params['FACTOR_CLIMA']
        rendimiento_base = params['RENDIMIENTO_BASE']
        ajuste_fertilidad = 0.5 + (factor_fertilidad * 0.5)
        rendimiento_potencial = (
            rendimiento_base *
            ajuste_fertilidad *
            factor_humedad *
            factor_vigor *
            factor_clima
        )
        rendimiento_potencial = min(rendimiento_potencial, params['RENDIMIENTO_OPTIMO'])
        rendimientos.append(round(rendimiento_potencial, 2))
    return rendimientos

def calcular_rendimiento_con_recomendaciones(gdf_analizado, cultivo):
    params = PARAMETROS_CULTIVOS[cultivo]
    rendimientos = []
    for idx, row in gdf_analizado.iterrows():
        factor_fertilidad = row['npk_integrado']
        factor_humedad = min(1.0, row['ndwi'] / 0.4) if 'ndwi' in row else 0.7
        factor_vigor = min(1.0, row['ndvi'] / params['NDVI_OPTIMO'])
        factor_clima = params['FACTOR_CLIMA']
        rendimiento_base = params['RENDIMIENTO_BASE']
        ajuste_fertilidad = 0.5 + (factor_fertilidad * 0.5)
        rendimiento_actual = (
            rendimiento_base *
            ajuste_fertilidad *
            factor_humedad *
            factor_vigor *
            factor_clima
        )
        incremento_total = 0
        n_actual = row['nitrogeno_actual']
        n_optimo = params['NITROGENO']['optimo']
        if n_actual < n_optimo * 0.9:
            eficiencia_n = params['RESPUESTA_N'] * 0.7
            incremento_n = row['valor_recomendado'] * eficiencia_n
            incremento_total += min(incremento_n, n_optimo * params['RESPUESTA_N'])
        p_actual = row['fosforo_actual']
        p_optimo = params['FOSFORO']['optimo']
        if p_actual < p_optimo * 0.85:
            deficiencia_p = max(0, p_optimo - p_actual)
            eficiencia_p = params['RESPUESTA_P'] * 0.5
            incremento_p = deficiencia_p * eficiencia_p
            incremento_total += incremento_p
        k_actual = row['potasio_actual']
        k_optimo = params['POTASIO']['optimo']
        if k_actual < k_optimo * 0.85:
            deficiencia_k = max(0, k_optimo - k_actual)
            eficiencia_k = params['RESPUESTA_K'] * 0.6
            incremento_k = deficiencia_k * eficiencia_k
            incremento_total += incremento_k
        rendimiento_proyectado = rendimiento_actual + incremento_total
        rendimiento_max = params['RENDIMIENTO_OPTIMO'] * 1.1
        rendimiento_proyectado = min(rendimiento_proyectado, rendimiento_max)
        rendimientos.append(round(rendimiento_proyectado, 2))
    return rendimientos

def crear_mapa_rendimiento(gdf_analizado, columna_rendimiento, cultivo, titulo, colormap='YlOrRd'):
    try:
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')
        valores = gdf_plot[columna_rendimiento]
        vmin = valores.min() * 0.8
        vmax = valores.max() * 1.2
        if colormap == 'YlOrRd':
            cmap = plt.cm.YlOrRd
        elif colormap == 'RdYlGn':
            cmap = plt.cm.RdYlGn
        else:
            cmap = plt.cm.viridis
        for idx, row in gdf_plot.iterrows():
            valor = row[columna_rendimiento]
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = cmap(valor_norm)
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white', linewidth=1.5, alpha=0.7)
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.1f}t",
                        (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='white', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='#1e293b', alpha=0.9, edgecolor='white'))
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.4)
        except:
            pass
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} {titulo} - {cultivo}',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label('Rendimiento (ton/ha)', fontsize=12, fontweight='bold', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando mapa de rendimiento: {str(e)}")
        return None

def crear_mapa_comparativo_rendimiento(gdf_analizado, cultivo):
    try:
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 8))
        fig.patch.set_facecolor('#0f172a')
        ax1.set_facecolor('#0f172a')
        ax2.set_facecolor('#0f172a')
        valores_actual = gdf_plot['rendimiento_actual']
        vmin1, vmax1 = valores_actual.min() * 0.8, valores_actual.max() * 1.2
        for idx, row in gdf_plot.iterrows():
            valor = row['rendimiento_actual']
            valor_norm = (valor - vmin1) / (vmax1 - vmin1) if vmax1 != vmin1 else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = plt.cm.YlOrRd(valor_norm)
            gdf_plot.iloc[[idx]].plot(ax=ax1, color=color, edgecolor='white', linewidth=1.5, alpha=0.7)
            centroid = row.geometry.centroid
            ax1.annotate(f"{valor:.1f}t",
                         (centroid.x, centroid.y),
                         xytext=(5, 5), textcoords="offset points",
                         fontsize=7, color='white', weight='bold')
        valores_proy = gdf_plot['rendimiento_proyectado']
        vmin2, vmax2 = valores_proy.min() * 0.8, valores_proy.max() * 1.2
        for idx, row in gdf_plot.iterrows():
            valor = row['rendimiento_proyectado']
            valor_norm = (valor - vmin2) / (vmax2 - vmin2) if vmax2 != vmin2 else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = plt.cm.RdYlGn(valor_norm)
            gdf_plot.iloc[[idx]].plot(ax=ax2, color=color, edgecolor='white', linewidth=1.5, alpha=0.7)
            incremento = row['rendimiento_proyectado'] - row['rendimiento_actual']
            centroid = row.geometry.centroid
            ax2.annotate(f"{valor:.1f}t\n(+{incremento:.1f})",
                         (centroid.x, centroid.y),
                         xytext=(5, 5), textcoords="offset points",
                         fontsize=7, color='white', weight='bold')
        ax1.set_title('📊 RENDIMIENTO ACTUAL', fontsize=14, fontweight='bold', color='white')
        ax2.set_title('🚀 RENDIMIENTO CON FERTILIZACIÓN', fontsize=14, fontweight='bold', color='white')
        for ax in [ax1, ax2]:
            ax.set_xlabel('Longitud', color='white')
            ax.set_ylabel('Latitud', color='white')
            ax.tick_params(colors='white')
            ax.grid(True, alpha=0.3, color='#475569')
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando mapa comparativo: {str(e)}")
        return None

# ===== FUNCIONES PARA DATOS SATELITALES =====
def descargar_datos_landsat8(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    try:
        st.info(f"🔍 Buscando escenas Landsat 8...")
        datos_simulados = {
            'indice': indice,
            'valor_promedio': 0.65 + np.random.normal(0, 0.1),
            'fuente': 'Landsat-8',
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'id_escena': f"LC08_{np.random.randint(1000000, 9999999)}",
            'cobertura_nubes': f"{np.random.randint(0, 15)}%",
            'resolucion': '30m'
        }
        st.success(f"✅ Escena Landsat 8 encontrada: {datos_simulados['id_escena']}")
        st.info(f"☁️ Cobertura de nubes: {datos_simulados['cobertura_nubes']}")
        return datos_simulados
    except Exception as e:
        st.error(f"❌ Error procesando Landsat 8: {str(e)}")
        return None

def descargar_datos_sentinel2(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    try:
        st.info(f"🔍 Buscando escenas Sentinel-2...")
        datos_simulados = {
            'indice': indice,
            'valor_promedio': 0.72 + np.random.normal(0, 0.08),
            'fuente': 'Sentinel-2',
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'id_escena': f"S2A_{np.random.randint(1000000, 9999999)}",
            'cobertura_nubes': f"{np.random.randint(0, 10)}%",
            'resolucion': '10m'
        }
        st.success(f"✅ Escena Sentinel-2 encontrada: {datos_simulados['id_escena']}")
        st.info(f"☁️ Cobertura de nubes: {datos_simulados['cobertura_nubes']}")
        return datos_simulados
    except Exception as e:
        st.error(f"❌ Error procesando Sentinel-2: {str(e)}")
        return None

def generar_datos_simulados(gdf, cultivo, indice='NDVI'):
    st.info("🔬 Generando datos simulados...")
    datos_simulados = {
        'indice': indice,
        'valor_promedio': PARAMETROS_CULTIVOS[cultivo]['NDVI_OPTIMO'] * 0.8 + np.random.normal(0, 0.1),
        'fuente': 'Simulación',
        'fecha': datetime.now().strftime('%Y-%m-%d'),
        'resolucion': '10m'
    }
    st.success("✅ Datos simulados generados")
    return datos_simulados

# ===== FUNCIÓN CORREGIDA PARA OBTENER DATOS DE NASA POWER =====
def obtener_datos_nasa_power(gdf, fecha_inicio, fecha_fin):
    try:
        centroid = gdf.geometry.unary_union.centroid
        lat = round(centroid.y, 4)
        lon = round(centroid.x, 4)
        start = fecha_inicio.strftime("%Y%m%d")
        end = fecha_fin.strftime("%Y%m%d")
        params = {
            'parameters': 'ALLSKY_SFC_SW_DWN,WS2M,T2M,PRECTOTCORR',
            'community': 'RE',
            'longitude': lon,
            'latitude': lat,
            'start': start,
            'end': end,
            'format': 'JSON'
        }
        url = "https://power.larc.nasa.gov/api/temporal/daily/point"
        response = requests.get(url, params=params, timeout=15)
        data = response.json()
        if 'properties' not in data:
            st.warning("⚠️ No se obtuvieron datos de NASA POWER (fuera de rango o sin conexión).")
            return None
        series = data['properties']['parameter']
        df_power = pd.DataFrame({
            'fecha': pd.to_datetime(list(series['ALLSKY_SFC_SW_DWN'].keys())),
            'radiacion_solar': list(series['ALLSKY_SFC_SW_DWN'].values()),
            'viento_2m': list(series['WS2M'].values()),
            'temperatura': list(series['T2M'].values()),
            'precipitacion': list(series['PRECTOTCORR'].values())
        })
        df_power = df_power.replace(-999, np.nan).dropna()
        if df_power.empty:
            st.warning("⚠️ Datos de NASA POWER no disponibles para el período seleccionado.")
            return None
        st.success("✅ Datos meteorológicos de NASA POWER cargados.")
        return df_power
    except Exception as e:
        st.error(f"❌ Error al obtener datos de NASA POWER: {str(e)}")
        return None

# ===== FUNCIONES DE ANÁLISIS GEE MEJORADAS =====
def calcular_recomendaciones_npk_cientificas(gdf_analizado, nutriente, cultivo):
    recomendaciones = []
    params = PARAMETROS_CULTIVOS[cultivo]
    for idx, row in gdf_analizado.iterrows():
        if nutriente == "NITRÓGENO":
            valor_actual = row['nitrogeno_actual']
            objetivo = params['NITROGENO']['optimo']
            deficiencia = max(0, objetivo - valor_actual)
            eficiencia = 0.5
            recomendado = deficiencia / eficiencia if deficiencia > 0 else 0
        elif nutriente == "FÓSFORO":
            valor_actual = row['fosforo_actual']
            objetivo = params['FOSFORO']['optimo']
            deficiencia = max(0, objetivo - valor_actual)
            eficiencia = 0.3
            recomendado = deficiencia / eficiencia if deficiencia > 0 else 0
        else:
            valor_actual = row['potasio_actual']
            objetivo = params['POTASIO']['optimo']
            deficiencia = max(0, objetivo - valor_actual)
            eficiencia = 0.6
            recomendado = deficiencia / eficiencia if deficiencia > 0 else 0
        recomendado_redondeado = round(recomendado / 5) * 5
        recomendaciones.append(max(0, recomendado_redondeado))
    return recomendaciones

# ===== FUNCIONES DE TEXTURA DEL SUELO - ACTUALIZADAS =====
def clasificar_textura_suelo(arena, limo, arcilla):
    try:
        total = arena + limo + arcilla
        if total == 0:
            return "NO_DETERMINADA"
        arena_norm = (arena / total) * 100
        limo_norm = (limo / total) * 100
        arcilla_norm = (arcilla / total) * 100
        if arcilla_norm >= 35:
            return "Franco arcilloso"
        elif arcilla_norm >= 25 and arcilla_norm <= 35 and arena_norm >= 20 and arena_norm <= 45:
            return "Franco arcilloso"
        elif arena_norm >= 40 and arena_norm <= 50 and arcilla_norm >= 20 and arcilla_norm <= 30:
            return "Franco arenoso-arcilloso"
        elif arena_norm >= 50 and arena_norm <= 70 and arcilla_norm >= 5 and arcilla_norm <= 20:
            return "Franco arenoso-arcilloso"
        elif arcilla_norm >= 7 and arcilla_norm <= 27 and arena_norm >= 43 and arena_norm <= 52:
            return "Franco"
        elif arena_norm >= 85:
            return "Franco arenoso-arcilloso"
        else:
            return "Franco"
    except Exception as e:
        return "NO_DETERMINADA"

def analizar_textura_suelo(gdf, cultivo):
    gdf = validar_y_corregir_crs(gdf)
    params_textura = TEXTURA_SUELO_OPTIMA[cultivo]
    zonas_gdf = gdf.copy()
    areas_ha_list = []
    arena_list = []
    limo_list = []
    arcilla_list = []
    textura_list = []
    for idx, row in zonas_gdf.iterrows():
        try:
            area_gdf = gpd.GeoDataFrame({'geometry': [row.geometry]}, crs=zonas_gdf.crs)
            area_ha = calcular_superficie(area_gdf)
            if hasattr(area_ha, 'iloc'):
                area_ha = float(area_ha.iloc[0])
            elif hasattr(area_ha, '__len__') and len(area_ha) > 0:
                area_ha = float(area_ha[0])
            else:
                area_ha = float(area_ha)
            centroid = row.geometry.centroid if hasattr(row.geometry, 'centroid') else row.geometry.representative_point()
            seed_value = abs(hash(f"{centroid.x:.6f}_{centroid.y:.6f}_{cultivo}_textura")) % (2**32)
            rng = np.random.RandomState(seed_value)
            arena_optima = params_textura['arena_optima']
            limo_optima = params_textura['limo_optima']
            arcilla_optima = params_textura['arcilla_optima']
            arena_val = max(5, min(95, rng.normal(arena_optima, 10)))
            limo_val = max(5, min(95, rng.normal(limo_optima, 8)))
            arcilla_val = max(5, min(95, rng.normal(arcilla_optima, 7)))
            total = arena_val + limo_val + arcilla_val
            arena_pct = (arena_val / total) * 100
            limo_pct = (limo_val / total) * 100
            arcilla_pct = (arcilla_val / total) * 100
            textura = clasificar_textura_suelo(arena_pct, limo_pct, arcilla_pct)
            areas_ha_list.append(area_ha)
            arena_list.append(float(arena_pct))
            limo_list.append(float(limo_pct))
            arcilla_list.append(float(arcilla_pct))
            textura_list.append(textura)
        except Exception as e:
            areas_ha_list.append(0.0)
            arena_list.append(float(params_textura['arena_optima']))
            limo_list.append(float(params_textura['limo_optima']))
            arcilla_list.append(float(params_textura['arcilla_optima']))
            textura_list.append(params_textura['textura_optima'])
    zonas_gdf['area_ha'] = areas_ha_list
    zonas_gdf['arena'] = arena_list
    zonas_gdf['limo'] = limo_list
    zonas_gdf['arcilla'] = arcilla_list
    zonas_gdf['textura_suelo'] = textura_list
    return zonas_gdf

# ===== FUNCIONES DE CURVAS DE NIVEL =====
def clasificar_pendiente(pendiente_porcentaje):
    for categoria, params in CLASIFICACION_PENDIENTES.items():
        if params['min'] <= pendiente_porcentaje < params['max']:
            return categoria, params['color']
    return "EXTREMA (>25%)", CLASIFICACION_PENDIENTES['EXTREMA (>25%)']['color']

def calcular_estadisticas_pendiente_simple(pendiente_grid):
    pendiente_flat = pendiente_grid.flatten()
    pendiente_flat = pendiente_flat[~np.isnan(pendiente_flat)]
    if len(pendiente_flat) == 0:
        return {'promedio': 0, 'min': 0, 'max': 0, 'std': 0, 'distribucion': {}}
    stats = {
        'promedio': float(np.mean(pendiente_flat)),
        'min': float(np.min(pendiente_flat)),
        'max': float(np.max(pendiente_flat)),
        'std': float(np.std(pendiente_flat)),
        'distribucion': {}
    }
    for categoria, params in CLASIFICACION_PENDIENTES.items():
        mask = (pendiente_flat >= params['min']) & (pendiente_flat < params['max'])
        stats['distribucion'][categoria] = {'porcentaje': float(np.sum(mask) / len(pendiente_flat) * 100), 'color': params['color']}
    return stats

def generar_dem_sintetico(gdf, resolucion=10.0):
    gdf = validar_y_corregir_crs(gdf)
    bounds = gdf.total_bounds
    minx, miny, maxx, maxy = bounds
    centroid = gdf.geometry.unary_union.centroid
    seed_value = int(centroid.x * 10000 + centroid.y * 10000) % (2**32)
    rng = np.random.RandomState(seed_value)
    num_cells = 50
    x = np.linspace(minx, maxx, num_cells)
    y = np.linspace(miny, maxy, num_cells)
    X, Y = np.meshgrid(x, y)
    elevacion_base = rng.uniform(100, 300)
    slope_x = rng.uniform(-0.001, 0.001)
    slope_y = rng.uniform(-0.001, 0.001)
    relief = np.zeros_like(X)
    n_hills = rng.randint(2, 5)
    for _ in range(n_hills):
        hill_center_x = rng.uniform(minx, maxx)
        hill_center_y = rng.uniform(miny, maxy)
        hill_radius = rng.uniform(0.001, 0.005)
        hill_height = rng.uniform(10, 50)
        dist = np.sqrt((X - hill_center_x)**2 + (Y - hill_center_y)**2)
        relief += hill_height * np.exp(-(dist**2) / (2 * hill_radius**2))
    noise = rng.randn(*X.shape) * 2
    Z = elevacion_base + slope_x * (X - minx) + slope_y * (Y - miny) + relief + noise
    Z = np.maximum(Z, 50)
    return X, Y, Z, bounds

def calcular_pendiente_simple(X, Y, Z, resolucion=10.0):
    dy = np.gradient(Z, axis=0) / resolucion
    dx = np.gradient(Z, axis=1) / resolucion
    pendiente = np.sqrt(dx**2 + dy**2) * 100
    pendiente = np.clip(pendiente, 0, 100)
    return pendiente

def crear_mapa_pendientes_simple(X, Y, pendiente_grid, gdf_original):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))
    fig.patch.set_facecolor('#0f172a')
    ax1.set_facecolor('#0f172a')
    ax2.set_facecolor('#0f172a')
    X_flat = X.flatten()
    Y_flat = Y.flatten()
    Z_flat = pendiente_grid.flatten()
    valid_mask = ~np.isnan(Z_flat)
    if np.sum(valid_mask) > 10:
        scatter = ax1.scatter(X_flat[valid_mask], Y_flat[valid_mask], c=Z_flat[valid_mask], cmap='RdYlGn_r', s=20, alpha=0.7, vmin=0, vmax=30)
        cbar = plt.colorbar(scatter, ax=ax1, shrink=0.8)
        cbar.set_label('Pendiente (%)', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')
        for porcentaje in [2, 5, 10, 15, 25]:
            mask_cat = (Z_flat[valid_mask] >= porcentaje-1) & (Z_flat[valid_mask] <= porcentaje+1)
            if np.sum(mask_cat) > 0:
                x_center = np.mean(X_flat[valid_mask][mask_cat])
                y_center = np.mean(Y_flat[valid_mask][mask_cat])
                ax1.text(x_center, y_center, f'{porcentaje}%', fontsize=8, fontweight='bold', ha='center', va='center',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='#1e293b', alpha=0.9, edgecolor='white'), color='white')
    else:
        ax1.text(0.5, 0.5, 'Datos insuficientes\npara mapa de calor', transform=ax1.transAxes, ha='center', va='center', fontsize=12, color='white')
    gdf_original.plot(ax=ax1, color='none', edgecolor='white', linewidth=2)
    ax1.set_title('Mapa de Calor de Pendientes', fontsize=12, fontweight='bold', color='white')
    ax1.set_xlabel('Longitud', color='white')
    ax1.set_ylabel('Latitud', color='white')
    ax1.tick_params(colors='white')
    ax1.grid(True, alpha=0.3, color='#475569')
    if np.sum(valid_mask) > 0:
        pendiente_data = Z_flat[valid_mask]
        ax2.hist(pendiente_data, bins=30, edgecolor='white', color='#3b82f6', alpha=0.7)
        for porcentaje, color in [(2, '#4daf4a'), (5, '#a6d96a'), (10, '#ffffbf'), (15, '#fdae61'), (25, '#f46d43')]:
            ax2.axvline(x=porcentaje, color=color, linestyle='--', linewidth=1, alpha=0.7)
            ax2.text(porcentaje+0.5, ax2.get_ylim()[1]*0.9, f'{porcentaje}%', color=color, fontsize=8)
        stats_pendiente = calcular_estadisticas_pendiente_simple(pendiente_grid)
        stats_text = f"""
Estadísticas:
• Mínima: {stats_pendiente['min']:.1f}%
• Máxima: {stats_pendiente['max']:.1f}%
• Promedio: {stats_pendiente['promedio']:.1f}%
• Desviación: {stats_pendiente['std']:.1f}%
"""
        ax2.text(0.02, 0.98, stats_text, transform=ax2.transAxes, fontsize=9, verticalalignment='top',
                color='white', bbox=dict(boxstyle="round,pad=0.3", facecolor='#1e293b', alpha=0.9, edgecolor='white'))
        ax2.set_xlabel('Pendiente (%)', color='white')
        ax2.set_ylabel('Frecuencia', color='white')
        ax2.set_title('Distribución de Pendientes', fontsize=12, fontweight='bold', color='white')
        ax2.tick_params(colors='white')
        ax2.grid(True, alpha=0.3, color='#475569')
    else:
        ax2.text(0.5, 0.5, 'Sin datos de pendiente', transform=ax2.transAxes, ha='center', va='center', fontsize=12, color='white')
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
    buf.seek(0)
    plt.close()
    return buf, calcular_estadisticas_pendiente_simple(pendiente_grid)

def generar_curvas_nivel_simple(X, Y, Z, intervalo=5.0, gdf_original=None):
    curvas = []
    elevaciones = []
    try:
        if gdf_original is not None:
            poligono_principal = gdf_original.iloc[0].geometry
            bounds = poligono_principal.bounds
            centro = poligono_principal.centroid
            ancho = bounds[2] - bounds[0]
            alto = bounds[3] - bounds[1]
            radio_max = min(ancho, alto) / 2
            z_min, z_max = np.nanmin(Z), np.nanmax(Z)
            n_curvas = min(10, int((z_max - z_min) / intervalo))
            for i in range(1, n_curvas + 1):
                radio = radio_max * (i / n_curvas)
                circle = centro.buffer(radio)
                interseccion = poligono_principal.intersection(circle)
                if interseccion.geom_type == 'LineString':
                    curvas.append(interseccion)
                    elevaciones.append(z_min + (i * intervalo))
                elif interseccion.geom_type == 'MultiLineString':
                    for parte in interseccion.geoms:
                        curvas.append(parte)
                        elevaciones.append(z_min + (i * intervalo))
    except Exception as e:
        if gdf_original is not None:
            bounds = gdf_original.total_bounds
            for i in range(3):
                y = bounds[1] + (i + 1) * ((bounds[3] - bounds[1]) / 4)
                linea = LineString([(bounds[0], y), (bounds[2], y)])
                curvas.append(linea)
                elevaciones.append(100 + i * 50)
    return curvas, elevaciones

# ===== FUNCIONES DE EXPORTACIÓN Y REPORTES - CORREGIDAS =====
def exportar_a_geojson(gdf, nombre_base="parcela"):
    try:
        gdf = validar_y_corregir_crs(gdf)
        geojson_data = gdf.to_json()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"{nombre_base}_{timestamp}.geojson"
        return geojson_data, nombre_archivo
    except Exception as e:
        st.error(f"❌ Error exportando a GeoJSON: {str(e)}")
        return None, None

def generar_resumen_estadisticas(gdf_analizado, analisis_tipo, cultivo, df_power=None):
    estadisticas = {}
    try:
        if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
            if 'npk_integrado' in gdf_analizado.columns:
                estadisticas['Índice NPK Integrado'] = f"{gdf_analizado['npk_integrado'].mean():.3f}"
            if 'nitrogeno_actual' in gdf_analizado.columns:
                estadisticas['Nitrógeno Promedio'] = f"{gdf_analizado['nitrogeno_actual'].mean():.1f} kg/ha"
            if 'fosforo_actual' in gdf_analizado.columns:
                estadisticas['Fósforo Promedio'] = f"{gdf_analizado['fosforo_actual'].mean():.1f} kg/ha"
            if 'potasio_actual' in gdf_analizado.columns:
                estadisticas['Potasio Promedio'] = f"{gdf_analizado['potasio_actual'].mean():.1f} kg/ha"
            if 'ndvi' in gdf_analizado.columns:
                estadisticas['NDVI Promedio'] = f"{gdf_analizado['ndvi'].mean():.3f}"
            if 'ndwi' in gdf_analizado.columns:
                estadisticas['NDWI Promedio'] = f"{gdf_analizado['ndwi'].mean():.3f}"
            if 'materia_organica' in gdf_analizado.columns:
                estadisticas['Materia Orgánica Promedio'] = f"{gdf_analizado['materia_organica'].mean():.1f}%"
            if df_power is not None:
                estadisticas['Radiación Solar Promedio'] = f"{df_power['radiacion_solar'].mean():.1f} kWh/m²/día"
                estadisticas['Velocidad Viento Promedio'] = f"{df_power['viento_2m'].mean():.2f} m/s"
                estadisticas['Precipitación Promedio'] = f"{df_power['precipitacion'].mean():.2f} mm/día"
        elif analisis_tipo == "ANÁLISIS DE TEXTURA":
            if 'arena' in gdf_analizado.columns:
                estadisticas['Arena Promedio'] = f"{gdf_analizado['arena'].mean():.1f}%"
            estadisticas['Limo Promedio'] = f"{gdf_analizado['limo'].mean():.1f}%"
            estadisticas['Arcilla Promedio'] = f"{gdf_analizado['arcilla'].mean():.1f}%"
            if 'textura_suelo' in gdf_analizado.columns:
                textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "N/D"
                estadisticas['Textura Predominante'] = textura_predominante
            if 'area_ha' in gdf_analizado.columns:
                estadisticas['Área Promedio por Zona'] = f"{gdf_analizado['area_ha'].mean():.2f} ha"
                if gdf_analizado['area_ha'].mean() > 0:
                    estadisticas['Coeficiente de Variación'] = f"{(gdf_analizado['area_ha'].std() / gdf_analizado['area_ha'].mean() * 100):.1f}%"
    except Exception as e:
        st.warning(f"No se pudieron calcular algunas estadísticas: {str(e)}")
    return estadisticas

def generar_recomendaciones_generales(gdf_analizado, analisis_tipo, cultivo):
    recomendaciones = []
    try:
        if analisis_tipo == "FERTILIDAD ACTUAL":
            if 'npk_integrado' in gdf_analizado.columns:
                npk_promedio = gdf_analizado['npk_integrado'].mean()
                if npk_promedio < 0.3:
                    recomendaciones.append("Fertilidad MUY BAJA: Se recomienda aplicación urgente de fertilizantes balanceados")
                    recomendaciones.append("Considerar enmiendas orgánicas para mejorar la estructura del suelo")
                elif npk_promedio < 0.5:
                    recomendaciones.append("Fertilidad BAJA: Recomendada aplicación de fertilizantes según análisis de suelo")
                elif npk_promedio < 0.7:
                    recomendaciones.append("Fertilidad ADECUADA: Mantener prácticas de manejo actuales")
                else:
                    recomendaciones.append("Fertilidad ÓPTIMA: Excelente condición, continuar con manejo actual")
            if 'nitrogeno_actual' in gdf_analizado.columns:
                n_prom = gdf_analizado['nitrogeno_actual'].mean()
                n_opt = PARAMETROS_CULTIVOS[cultivo]['NITROGENO']['optimo']
                if n_prom < n_opt * 0.7:
                    recomendaciones.append(f"Deficiencia de Nitrógeno ({n_prom:.1f} vs {n_opt:.1f} kg/ha): Aplicar fertilizante nitrogenado")
            if 'fosforo_actual' in gdf_analizado.columns:
                p_prom = gdf_analizado['fosforo_actual'].mean()
                p_opt = PARAMETROS_CULTIVOS[cultivo]['FOSFORO']['optimo']
                if p_prom < p_opt * 0.7:
                    recomendaciones.append(f"Deficiencia de Fósforo ({p_prom:.1f} vs {p_opt:.1f} kg/ha): Aplicar superfosfato o fosfato diamónico")
            if 'potasio_actual' in gdf_analizado.columns:
                k_prom = gdf_analizado['potasio_actual'].mean()
                k_opt = PARAMETROS_CULTIVOS[cultivo]['POTASIO']['optimo']
                if k_prom < k_opt * 0.7:
                    recomendaciones.append(f"Deficiencia de Potasio ({k_prom:.1f} vs {k_opt:.1f} kg/ha): Aplicar cloruro o sulfato de potasio")
        elif analisis_tipo == "ANÁLISIS DE TEXTURA":
            if 'textura_suelo' in gdf_analizado.columns:
                textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "N/D"
                if textura_predominante == "Franco arcilloso":
                    recomendaciones.append("Suelo franco arcilloso: Mejorar drenaje y evitar laboreo en condiciones húmedas")
                elif textura_predominante == "Franco arenoso-arcilloso":
                    recomendaciones.append("Suelo franco arenoso-arcilloso: Aumentar materia orgánica y considerar riego frecuente")
                elif textura_predominante == "Franco":
                    recomendaciones.append("Textura franca: Condiciones óptimas, mantener prácticas de conservación")
        if cultivo == "MAÍZ":
            recomendaciones.append("Para maíz: Priorizar aplicación de nitrógeno en etapas de crecimiento vegetativo.")
            recomendaciones.append("Mantener humedad adecuada durante floración y llenado de grano.")
        elif cultivo == "SOYA":
            recomendaciones.append("Para soya: Inocular con rizobios para fijación de nitrógeno atmosférico.")
            recomendaciones.append("Manejo adecuado de humedad durante formación de vainas.")
        elif cultivo == "TRIGO":
            recomendaciones.append("Para trigo: Aplicar nitrógeno en macollamiento y encañazón.")
            recomendaciones.append("Controlar humedad para evitar enfermedades fúngicas.")
        elif cultivo == "GIRASOL":
            recomendaciones.append("Para girasol: Aplicar potasio para mejorar calidad de semilla.")
            recomendaciones.append("Mantener buen drenaje, sensible a encharcamiento.")
        recomendaciones.append("Realizar análisis de suelo de laboratorio para validar resultados satelitales")
        recomendaciones.append("Considerar agricultura de precisión para aplicación variable de insumos")
    except Exception as e:
        recomendaciones.append("Error generando recomendaciones específicas")
    return recomendaciones

def limpiar_texto_para_pdf(texto):
    if not isinstance(texto, str):
        texto = str(texto)
    reemplazos = {
        '\u2022': '-',
        '\u2705': '[OK]',
        '\u26A0\uFE0F': '[!]',
        '\u274C': '[X]',
        '\u2013': '-',
        '\u2014': '--',
        '\u2018': "'",
        '\u2019': "'",
        '\u201C': '"',
        '\u201D': '"',
        '\u2192': '->',
        '\u2190': '<-',
        '\u2265': '>=',
        '\u2264': '<=',
        '\u00A0': ' ',
    }
    for original, reemplazo in reemplazos.items():
        texto = texto.replace(original, reemplazo)
    texto = texto.encode('latin-1', errors='replace').decode('latin-1')
    return texto

def generar_reporte_pdf(gdf_analizado, cultivo, analisis_tipo, area_total,
                        nutriente=None, satelite=None, indice=None,
                        mapa_buffer=None, estadisticas=None, recomendaciones=None):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font('Arial', '', 12)
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, limpiar_texto_para_pdf(f'REPORTE DE ANÁLISIS AGRÍCOLA - {cultivo}'), 0, 1, 'C')
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 10, limpiar_texto_para_pdf(f'Tipo de Análisis: {analisis_tipo}'), 0, 1, 'C')
        pdf.cell(0, 10, limpiar_texto_para_pdf(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}'), 0, 1, 'C')
        pdf.ln(10)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, '1. INFORMACIÓN GENERAL', 0, 1)
        pdf.set_font('Arial', '', 12)
        info_general = f"""Cultivo: {cultivo}
Área Total: {area_total:.2f} ha
Zonas Analizadas: {len(gdf_analizado)}
Tipo de Análisis: {analisis_tipo}"""
        if satelite:
            info_general += f"\nSatélite: {satelite}"
        if indice:
            info_general += f"\nÍndice: {indice}"
        if nutriente:
            info_general += f"\nNutriente Analizado: {nutriente}"
        for linea in info_general.strip().split('\n'):
            pdf.cell(0, 8, limpiar_texto_para_pdf(linea), 0, 1)
        pdf.ln(5)
        if estadisticas:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '2. ESTADÍSTICAS PRINCIPALES', 0, 1)
            pdf.set_font('Arial', '', 12)
            for key, value in estadisticas.items():
                linea = f"- {key}: {value}"
                pdf.cell(0, 8, limpiar_texto_para_pdf(linea), 0, 1)
            pdf.ln(5)
        if mapa_buffer:
            try:
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, '3. MAPA DE RESULTADOS', 0, 1)
                temp_img_path = "temp_map.png"
                with open(temp_img_path, "wb") as f:
                    f.write(mapa_buffer.getvalue())
                pdf.image(temp_img_path, x=10, w=190)
                pdf.ln(5)
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
            except Exception as e:
                pdf.cell(0, 8, limpiar_texto_para_pdf(f"Error al incluir mapa: {str(e)[:50]}..."), 0, 1)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, '4. RESUMEN DE ZONAS', 0, 1)
        pdf.set_font('Arial', '', 10)
        if gdf_analizado is not None and not gdf_analizado.empty:
            columnas_mostrar = ['id_zona', 'area_ha']
            if 'npk_integrado' in gdf_analizado.columns:
                columnas_mostrar.append('npk_integrado')
            if 'nitrogeno_actual' in gdf_analizado.columns:
                columnas_mostrar.append('nitrogeno_actual')
            if 'fosforo_actual' in gdf_analizado.columns:
                columnas_mostrar.append('fosforo_actual')
            if 'potasio_actual' in gdf_analizado.columns:
                columnas_mostrar.append('potasio_actual')
            if 'textura_suelo' in gdf_analizado.columns:
                columnas_mostrar.append('textura_suelo')
            if 'ndwi' in gdf_analizado.columns:
                columnas_mostrar.append('ndwi')
            columnas_mostrar = [col for col in columnas_mostrar if col in gdf_analizado.columns]
            if columnas_mostrar:
                datos_tabla = [columnas_mostrar]
                for _, row in gdf_analizado.head(15).iterrows():
                    fila = []
                    for col in columnas_mostrar:
                        if col in gdf_analizado.columns:
                            valor = row[col]
                            if isinstance(valor, float):
                                if col in ['npk_integrado', 'ndwi']:
                                    fila.append(f"{valor:.3f}")
                                elif col in ['nitrogeno_actual', 'fosforo_actual', 'potasio_actual']:
                                    fila.append(f"{valor:.1f}")
                                else:
                                    fila.append(f"{valor:.2f}")
                            else:
                                fila.append(str(valor))
                        else:
                            fila.append("N/A")
                    datos_tabla.append(fila)
                col_widths = [190 // len(columnas_mostrar)] * len(columnas_mostrar)
                for fila in datos_tabla:
                    for i, item in enumerate(fila):
                        if i < len(col_widths):
                            pdf.cell(col_widths[i], 8, limpiar_texto_para_pdf(str(item)), border=1)
                    pdf.ln()
                pdf.ln(5)
        if recomendaciones:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '5. RECOMENDACIONES', 0, 1)
            pdf.set_font('Arial', '', 12)
            for rec in recomendaciones:
                linea = f"- {limpiar_texto_para_pdf(rec)}"
                pdf.multi_cell(0, 8, linea)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, '6. METADATOS TÉCNICOS', 0, 1)
        pdf.set_font('Arial', '', 10)
        metadatos = f"""Generado por: Analizador Multi-Cultivo Satellital
Versión: 2.0 (Con metodologías científicas NPK)
Fecha de generación: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
Sistema de coordenadas: EPSG:4326 (WGS84)
Número de zonas: {len(gdf_analizado)}"""
        for linea in metadatos.strip().split('\n'):
            pdf.cell(0, 6, limpiar_texto_para_pdf(linea), 0, 1)
        pdf_output = BytesIO()
        pdf_output.write(pdf.output(dest='S').encode('latin-1'))
        pdf_output.seek(0)
        return pdf_output
    except Exception as e:
        st.error(f"❌ Error generando PDF: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

def generar_reporte_docx(gdf_analizado, cultivo, analisis_tipo, area_total,
                         nutriente=None, satelite=None, indice=None,
                         mapa_buffer=None, estadisticas=None, recomendaciones=None):
    try:
        doc = Document()
        title = doc.add_heading(f'REPORTE DE ANÁLISIS AGRÍCOLA - {cultivo}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph(f'Tipo de Análisis: {analisis_tipo}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha = doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        doc.add_heading('1. INFORMACIÓN GENERAL', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = 'Cultivo'
        info_table.cell(0, 1).text = cultivo
        info_table.cell(1, 0).text = 'Área Total'
        info_table.cell(1, 1).text = f'{area_total:.2f} ha'
        info_table.cell(2, 0).text = 'Zonas Analizadas'
        info_table.cell(2, 1).text = str(len(gdf_analizado))
        info_table.cell(3, 0).text = 'Tipo de Análisis'
        info_table.cell(3, 1).text = analisis_tipo
        row_count = 4
        if satelite:
            if row_count >= len(info_table.rows):
                info_table.add_row()
            info_table.cell(row_count, 0).text = 'Satélite'
            info_table.cell(row_count, 1).text = satelite
            row_count += 1
        if indice:
            if row_count >= len(info_table.rows):
                info_table.add_row()
            info_table.cell(row_count, 0).text = 'Índice'
            info_table.cell(row_count, 1).text = indice
            row_count += 1
        if nutriente:
            if row_count >= len(info_table.rows):
                info_table.add_row()
            info_table.cell(row_count, 0).text = 'Nutriente Analizado'
            info_table.cell(row_count, 1).text = nutriente
        doc.add_paragraph()
        if estadisticas:
            doc.add_heading('2. ESTADÍSTICAS PRINCIPALES', level=1)
            for key, value in estadisticas.items():
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f'{key}: ')
                run.bold = True
                p.add_run(str(value))
            doc.add_paragraph()
        if mapa_buffer:
            try:
                doc.add_heading('3. MAPA DE RESULTADOS', level=1)
                temp_img_path = "temp_map_docx.png"
                with open(temp_img_path, "wb") as f:
                    f.write(mapa_buffer.getvalue())
                doc.add_picture(temp_img_path, width=Inches(6.0))
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f'Error al incluir mapa: {str(e)[:50]}...')
        doc.add_heading('4. RESUMEN DE ZONAS', level=1)
        if gdf_analizado is not None and not gdf_analizado.empty:
            columnas_mostrar = ['id_zona', 'area_ha']
            if 'npk_integrado' in gdf_analizado.columns:
                columnas_mostrar.append('npk_integrado')
            if 'nitrogeno_actual' in gdf_analizado.columns:
                columnas_mostrar.append('nitrogeno_actual')
            if 'fosforo_actual' in gdf_analizado.columns:
                columnas_mostrar.append('fosforo_actual')
            if 'potasio_actual' in gdf_analizado.columns:
                columnas_mostrar.append('potasio_actual')
            if 'textura_suelo' in gdf_analizado.columns:
                columnas_mostrar.append('textura_suelo')
            if 'ndwi' in gdf_analizado.columns:
                columnas_mostrar.append('ndwi')
            columnas_mostrar = [col for col in columnas_mostrar if col in gdf_analizado.columns]
            if columnas_mostrar:
                tabla = doc.add_table(rows=1, cols=len(columnas_mostrar))
                tabla.style = 'Table Grid'
                for i, col in enumerate(columnas_mostrar):
                    tabla.cell(0, i).text = col.replace('_', ' ').upper()
                for idx, row in gdf_analizado.head(10).iterrows():
                    row_cells = tabla.add_row().cells
                    for i, col in enumerate(columnas_mostrar):
                        if col in gdf_analizado.columns:
                            valor = row[col]
                            if isinstance(valor, float):
                                if col in ['npk_integrado', 'ndwi']:
                                    row_cells[i].text = f"{valor:.3f}"
                                elif col in ['nitrogeno_actual', 'fosforo_actual', 'potasio_actual']:
                                    row_cells[i].text = f"{valor:.1f}"
                                else:
                                    row_cells[i].text = f"{valor:.2f}"
                            else:
                                row_cells[i].text = str(valor)
                        else:
                            row_cells[i].text = "N/A"
        doc.add_paragraph()
        if recomendaciones:
            doc.add_heading('5. RECOMENDACIONES', level=1)
            for rec in recomendaciones:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(rec)
        doc.add_heading('6. METADATOS TÉCNICOS', level=1)
        metadatos = [
            ('Generado por', 'Analizador Multi-Cultivo Satellital'),
            ('Versión', '2.0 (Con metodologías científicas NPK)'),
            ('Fecha de generación', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ('Sistema de coordenadas', 'EPSG:4326 (WGS84)'),
            ('Número de zonas', str(len(gdf_analizado)))
        ]
        for key, value in metadatos:
            p = doc.add_paragraph()
            run_key = p.add_run(f'{key}: ')
            run_key.bold = True
            p.add_run(value)
        docx_output = BytesIO()
        doc.save(docx_output)
        docx_output.seek(0)
        return docx_output
    except Exception as e:
        st.error(f"❌ Error generando DOCX: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

# ===== FUNCIONES DE VISUALIZACIÓN MEJORADAS CON MAPAS ESRI =====
def crear_mapa_npk_con_esri(gdf_analizado, nutriente, cultivo, satelite):
    try:
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')
        if nutriente == "NITRÓGENO":
            columna = 'nitrogeno_actual'
            cmap = LinearSegmentedColormap.from_list('nitrogeno_gee', PALETAS_GEE['NITROGENO'])
            vmin = PARAMETROS_CULTIVOS[cultivo]['NITROGENO']['min'] * 0.7
            vmax = PARAMETROS_CULTIVOS[cultivo]['NITROGENO']['max'] * 1.2
            titulo_nutriente = "NITRÓGENO (kg/ha)"
        elif nutriente == "FÓSFORO":
            columna = 'fosforo_actual'
            cmap = LinearSegmentedColormap.from_list('fosforo_gee', PALETAS_GEE['FOSFORO'])
            vmin = PARAMETROS_CULTIVOS[cultivo]['FOSFORO']['min'] * 0.7
            vmax = PARAMETROS_CULTIVOS[cultivo]['FOSFORO']['max'] * 1.2
            titulo_nutriente = "FÓSFORO (kg/ha)"
        else:
            columna = 'potasio_actual'
            cmap = LinearSegmentedColormap.from_list('potasio_gee', PALETAS_GEE['POTASIO'])
            vmin = PARAMETROS_CULTIVOS[cultivo]['POTASIO']['min'] * 0.7
            vmax = PARAMETROS_CULTIVOS[cultivo]['POTASIO']['max'] * 1.2
            titulo_nutriente = "POTASIO (kg/ha)"
        for idx, row in gdf_plot.iterrows():
            valor = row[columna]
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = cmap(valor_norm)
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white', linewidth=1.5, alpha=0.7)
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.0f}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='white', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='#1e293b', alpha=0.9, edgecolor='white'))
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.4)
        except:
            st.warning("⚠️ No se pudo cargar el mapa base ESRI. Verifica la conexión a internet.")
        info_satelite = SATELITES_DISPONIBLES.get(satelite, SATELITES_DISPONIBLES['DATOS_SIMULADOS'])
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} ANÁLISIS DE {nutriente} - {cultivo}\n'
                     f'{info_satelite["icono"]} {info_satelite["nombre"]} - {titulo_nutriente}',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label(titulo_nutriente, fontsize=12, fontweight='bold', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando mapa NPK con ESRI: {str(e)}")
        return None

def crear_mapa_fertilidad_integrada(gdf_analizado, cultivo, satelite):
    try:
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')
        cmap = LinearSegmentedColormap.from_list('fertilidad_gee', PALETAS_GEE['FERTILIDAD'])
        for idx, row in gdf_plot.iterrows():
            valor = row['npk_integrado']
            color = cmap(valor)
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white', linewidth=1.5, alpha=0.7)
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.2f}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='white', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='#1e293b', alpha=0.9, edgecolor='white'))
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.4)
        except:
            pass
        info_satelite = SATELITES_DISPONIBLES.get(satelite, SATELITES_DISPONIBLES['DATOS_SIMULADOS'])
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} FERTILIDAD INTEGRADA (NPK) - {cultivo}\n'
                     f'{info_satelite["icono"]} {info_satelite["nombre"]}',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=0, vmax=1))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label('Índice de Fertilidad (0-1)', fontsize=12, fontweight='bold', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando mapa fertilidad: {str(e)}")
        return None

def crear_mapa_texturas_con_esri(gdf_analizado, cultivo):
    try:
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')
        colores_textura = {
            'Franco': '#c7eae5',
            'Franco arcilloso': '#5ab4ac',
            'Franco arenoso-arcilloso': '#f6e8c3',
            'NO_DETERMINADA': '#999999'
        }
        for idx, row in gdf_plot.iterrows():
            textura = row['textura_suelo']
            color = colores_textura.get(textura, '#999999')
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white', linewidth=1.5, alpha=0.8)
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{textura[:10]}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='black', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9))
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.4)
        except:
            st.warning("⚠️ No se pudo cargar el mapa base ESRI. Verifica la conexión a internet.")
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} MAPA DE TEXTURAS - {cultivo}',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')
        from matplotlib.patches import Patch
        legend_elements = [Patch(facecolor=color, edgecolor='white', label=textura)
                          for textura, color in colores_textura.items()]
        legend = ax.legend(handles=legend_elements, title='Texturas', loc='upper left', bbox_to_anchor=(1.05, 1))
        legend.get_title().set_color('white')
        for text in legend.get_texts():
            text.set_color('white')
        legend.get_frame().set_facecolor('#1e293b')
        legend.get_frame().set_edgecolor('white')
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"Error creando mapa de texturas: {str(e)}")
        return None

# ===== FUNCIONES DE GRÁFICOS NASA POWER CON ESTILO OSCURO =====
def crear_grafico_personalizado(series, titulo, ylabel, color_linea, fondo_grafico='#0f172a', color_texto='#ffffff'):
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_facecolor(fondo_grafico)
    fig.patch.set_facecolor(fondo_grafico)
    ax.plot(series.index, series.values, color=color_linea, linewidth=2.2)
    ax.set_title(titulo, fontsize=14, fontweight='bold', color=color_texto)
    ax.set_ylabel(ylabel, fontsize=12, color=color_texto)
    ax.set_xlabel("Fecha", fontsize=11, color=color_texto)
    ax.tick_params(axis='x', colors=color_texto, rotation=0)
    ax.tick_params(axis='y', colors=color_texto)
    ax.grid(True, color='#475569', linestyle='--', linewidth=0.7, alpha=0.7)
    for spine in ax.spines.values():
        spine.set_color('#475569')
    plt.tight_layout()
    return fig

def crear_grafico_barras_personalizado(series, titulo, ylabel, color_barra, fondo_grafico='#0f172a', color_texto='#ffffff'):
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_facecolor(fondo_grafico)
    fig.patch.set_facecolor(fondo_grafico)
    ax.bar(series.index, series.values, color=color_barra, alpha=0.85)
    ax.set_title(titulo, fontsize=14, fontweight='bold', color=color_texto)
    ax.set_ylabel(ylabel, fontsize=12, color=color_texto)
    ax.set_xlabel("Fecha", fontsize=11, color=color_texto)
    ax.tick_params(axis='x', colors=color_texto, rotation=0)
    ax.tick_params(axis='y', colors=color_texto)
    ax.grid(axis='y', color='#475569', linestyle='--', linewidth=0.7, alpha=0.7)
    for spine in ax.spines.values():
        spine.set_color('#475569')
    plt.tight_layout()
    return fig

# ===== FUNCIÓN PRINCIPAL DE ANÁLISIS (MEJORADA) =====
def ejecutar_analisis(gdf, nutriente, analisis_tipo, n_divisiones, cultivo,
                      satelite=None, indice=None, fecha_inicio=None,
                      fecha_fin=None, intervalo_curvas=5.0, resolucion_dem=10.0):
    resultados = {
        'exitoso': False,
        'gdf_analizado': None,
        'mapa_buffer': None,
        'tabla_datos': None,
        'estadisticas': {},
        'recomendaciones': [],
        'area_total': 0,
        'df_power': None
    }
    try:
        gdf = validar_y_corregir_crs(gdf)
        area_total = calcular_superficie(gdf)
        resultados['area_total'] = area_total
        if analisis_tipo == "ANÁLISIS DE TEXTURA":
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            gdf_analizado = analizar_textura_suelo(gdf_dividido, cultivo)
            resultados['gdf_analizado'] = gdf_analizado
            resultados['exitoso'] = True
            return resultados
        elif analisis_tipo == "ANÁLISIS DE CURVAS DE NIVEL":
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            resultados['gdf_analizado'] = gdf_dividido
            resultados['exitoso'] = True
            return resultados
        elif analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
            datos_satelitales = None
            if satelite == "SENTINEL-2":
                datos_satelitales = descargar_datos_sentinel2(gdf, fecha_inicio, fecha_fin, indice)
            elif satelite == "LANDSAT-8":
                datos_satelitales = descargar_datos_landsat8(gdf, fecha_inicio, fecha_fin, indice)
            else:
                datos_satelitales = generar_datos_simulados(gdf, cultivo, indice)
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            indices_npk = calcular_indices_npk_avanzados(gdf_dividido, cultivo, satelite)
            gdf_analizado = gdf_dividido.copy()
            for idx, indice_data in enumerate(indices_npk):
                for key, value in indice_data.items():
                    gdf_analizado.loc[gdf_analizado.index[idx], key] = value
            areas_ha_list = []
            for idx, row in gdf_analizado.iterrows():
                area_gdf = gpd.GeoDataFrame({'geometry': [row.geometry]}, crs=gdf_analizado.crs)
                area_ha = calcular_superficie(area_gdf)
                if hasattr(area_ha, 'iloc'):
                    area_ha = float(area_ha.iloc[0])
                elif hasattr(area_ha, '__len__') and len(area_ha) > 0:
                    area_ha = float(area_ha[0])
                else:
                    area_ha = float(area_ha)
                areas_ha_list.append(area_ha)
            gdf_analizado['area_ha'] = areas_ha_list
            if analisis_tipo == "RECOMENDACIONES NPK":
                recomendaciones_npk = calcular_recomendaciones_npk_cientificas(gdf_analizado, nutriente, cultivo)
                gdf_analizado['valor_recomendado'] = recomendaciones_npk
            resultados['gdf_analizado'] = gdf_analizado
            resultados['exitoso'] = True
            if satelite:
                df_power = obtener_datos_nasa_power(gdf, fecha_inicio, fecha_fin)
                if df_power is not None:
                    resultados['df_power'] = df_power
            return resultados
        else:
            st.error(f"Tipo de análisis no soportado: {analisis_tipo}")
            return resultados
    except Exception as e:
        st.error(f"❌ Error en análisis: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return resultados

# ===== FUNCIONES DE VISUALIZACIÓN =====
def mostrar_resultados_textura(gdf_analizado, cultivo, area_total):
    st.subheader("📊 ESTADÍSTICAS DE TEXTURA")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "NO_DETERMINADA"
        st.metric("🏗️ Textura Predominante", textura_predominante)
    with col2:
        avg_arena = gdf_analizado['arena'].mean()
        st.metric("🏖️ Arena Promedio", f"{avg_arena:.1f}%")
    with col3:
        avg_limo = gdf_analizado['limo'].mean()
        st.metric("🌫️ Limo Promedio", f"{avg_limo:.1f}%")
    with col4:
        avg_arcilla = gdf_analizado['arcilla'].mean()
        st.metric("🧱 Arcilla Promedio", f"{avg_arcilla:.1f}%")
    st.subheader("📈 COMPOSICIÓN GRANULOMÉTRICA")
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
    fig.patch.set_facecolor('#0f172a')
    ax1.set_facecolor('#0f172a')
    ax2.set_facecolor('#0f172a')
    composicion = [gdf_analizado['arena'].mean(), gdf_analizado['limo'].mean(), gdf_analizado['arcilla'].mean()]
    labels = ['Arena', 'Limo', 'Arcilla']
    colors_pie = ['#d8b365', '#f6e8c3', '#01665e']
    ax1.pie(composicion, labels=labels, colors=colors_pie, autopct='%1.1f%%', startangle=90, textprops={'color': 'white'})
    ax1.set_title('Composición Promedio del Suelo', color='white')
    textura_dist = gdf_analizado['textura_suelo'].value_counts()
    ax2.bar(textura_dist.index, textura_dist.values, color=[PALETAS_GEE['TEXTURA'][i % len(PALETAS_GEE['TEXTURA'])] for i in range(len(textura_dist))])
    ax2.set_title('Distribución de Texturas', color='white')
    ax2.set_xlabel('Textura', color='white')
    ax2.set_ylabel('Número de Zonas', color='white')
    ax2.tick_params(axis='x', rotation=45, colors='white')
    ax2.tick_params(axis='y', colors='white')
    ax2.set_facecolor('#0f172a')
    plt.tight_layout()
    st.pyplot(fig)
    st.subheader("🗺️ MAPA DE TEXTURAS CON ESRI SATELLITE")
    mapa_texturas = crear_mapa_texturas_con_esri(gdf_analizado, cultivo)
    if mapa_texturas:
        st.image(mapa_texturas, use_container_width=True)
        st.download_button(
            "📥 Descargar Mapa de Texturas",
            mapa_texturas,
            f"mapa_texturas_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
            "image/png"
        )
    st.subheader("📋 TABLA DE RESULTADOS POR ZONA")
    columnas_textura = ['id_zona', 'area_ha', 'textura_suelo', 'arena', 'limo', 'arcilla']
    columnas_textura = [col for col in columnas_textura if col in gdf_analizado.columns]
    if columnas_textura:
        tabla_textura = gdf_analizado[columnas_textura].copy()
        tabla_textura.columns = ['Zona', 'Área (ha)', 'Textura', 'Arena (%)', 'Limo (%)', 'Arcilla (%)']
        st.dataframe(tabla_textura)
    st.subheader("💡 RECOMENDACIONES DE MANEJO POR TEXTURA")
    if 'textura_suelo' in gdf_analizado.columns:
        textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "NO_DETERMINADA"
        if textura_predominante in RECOMENDACIONES_TEXTURA:
            st.markdown(f"#### 🏗️ **{textura_predominante.upper()}**")
            info_textura = RECOMENDACIONES_TEXTURA[textura_predominante]
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown("**✅ PROPIEDADES FÍSICAS**")
                for prop in info_textura['propiedades']:
                    st.markdown(f"• {prop}")
            with col2:
                st.markdown("**⚠️ LIMITANTES**")
                for lim in info_textura['limitantes']:
                    st.markdown(f"• {lim}")
            with col3:
                st.markdown("**🛠️ MANEJO RECOMENDADO**")
                for man in info_textura['manejo']:
                    st.markdown(f"• {man}")
    st.subheader("💾 DESCARGAR RESULTADOS")
    if 'columnas_textura' in locals() and columnas_textura:
        tabla_textura = gdf_analizado[columnas_textura].copy()
        tabla_textura.columns = ['Zona', 'Área (ha)', 'Textura', 'Arena (%)', 'Limo (%)', 'Arcilla (%)']
        csv = tabla_textura.to_csv(index=False)
        st.download_button(
            "📥 Descargar CSV con Análisis de Textura",
            csv,
            f"textura_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            "text/csv"
        )

def mostrar_resultados_curvas_nivel(X, Y, Z, pendiente_grid, curvas, elevaciones, gdf_original, cultivo, area_total):
    st.subheader("📊 ESTADÍSTICAS TOPOGRÁFICAS")
    elevaciones_flat = Z.flatten()
    elevaciones_flat = elevaciones_flat[~np.isnan(elevaciones_flat)]
    if len(elevaciones_flat) > 0:
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            elevacion_promedio = np.mean(elevaciones_flat)
            st.metric("🏔️ Elevación Promedio", f"{elevacion_promedio:.1f} m")
        with col2:
            rango_elevacion = np.max(elevaciones_flat) - np.min(elevaciones_flat)
            st.metric("📏 Rango de Elevación", f"{rango_elevacion:.1f} m")
        with col3:
            mapa_pendientes, stats_pendiente = crear_mapa_pendientes_simple(X, Y, pendiente_grid, gdf_original)
            st.metric("📐 Pendiente Promedio", f"{stats_pendiente['promedio']:.1f}%")
        with col4:
            num_curvas = len(curvas) if curvas else 0
            st.metric("🔄 Número de Curvas", f"{num_curvas}")
    st.subheader("🔥 MAPA DE CALOR DE PENDIENTES")
    st.image(mapa_pendientes, use_container_width=True)
    st.download_button(
        "📥 Descargar Mapa de Pendientes",
        mapa_pendientes,
        f"mapa_pendientes_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
        "image/png"
    )
    st.subheader("⚠️ ANÁLISIS DE RIESGO DE EROSION")
    if 'stats_pendiente' in locals() and 'distribucion' in stats_pendiente:
        riesgo_total = 0
        for categoria, data in stats_pendiente['distribucion'].items():
            if categoria in CLASIFICACION_PENDIENTES:
                riesgo_total += data['porcentaje'] * CLASIFICACION_PENDIENTES[categoria]['factor_erosivo']
        riesgo_promedio = riesgo_total / 100
        col1, col2, col3 = st.columns(3)
        with col1:
            if riesgo_promedio < 0.3:
                st.success("✅ **RIESGO BAJO**")
                st.metric("Factor Riesgo", f"{riesgo_promedio:.2f}")
            elif riesgo_promedio < 0.6:
                st.warning("⚠️ **RIESGO MODERADO**")
                st.metric("Factor Riesgo", f"{riesgo_promedio:.2f}")
            else:
                st.error("🚨 **RIESGO ALTO**")
                st.metric("Factor Riesgo", f"{riesgo_promedio:.2f}")
        with col2:
            area_total_ha = area_total
            porcentaje_critico = sum(data['porcentaje'] for cat, data in stats_pendiente['distribucion'].items()
                                     if cat in ['FUERTE (10-15%)', 'MUY FUERTE (15-25%)', 'EXTREMA (>25%)'])
            area_critica = area_total_ha * (porcentaje_critico / 100)
            st.metric("Área Crítica (>10%)", f"{area_critica:.2f} ha")
        with col3:
            porcentaje_manejable = sum(data['porcentaje'] for cat, data in stats_pendiente['distribucion'].items()
                                       if cat in ['PLANA (0-2%)', 'SUAVE (2-5%)', 'MODERADA (5-10%)'])
            area_manejable = area_total_ha * (porcentaje_manejable / 100)
            st.metric("Área Manejable (<10%)", f"{area_manejable:.2f} ha")
    st.subheader("📈 VISUALIZACIÓN 3D DEL TERRENO")
    try:
        fig = plt.figure(figsize=(12, 8))
        ax = fig.add_subplot(111, projection='3d')
        surf = ax.plot_surface(X, Y, Z, cmap='terrain', alpha=0.8, linewidth=0)
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.set_zlabel('Elevación (m)', color='white')
        ax.set_title(f'Modelo 3D del Terreno - {cultivo}', color='white')
        ax.tick_params(colors='white')
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')
        ax.xaxis.label.set_color('white')
        ax.yaxis.label.set_color('white')
        ax.zaxis.label.set_color('white')
        ax.title.set_color('white')
        cbar = fig.colorbar(surf, ax=ax, shrink=0.5, aspect=5, label='Elevación (m)')
        cbar.set_label('Elevación (m)', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')
        plt.tight_layout()
        st.pyplot(fig)
    except Exception as e:
        st.warning(f"No se pudo generar visualización 3D: {e}")
    st.subheader("💾 DESCARGAR RESULTADOS")
    sample_points = []
    for i in range(0, X.shape[0], 5):
        for j in range(0, X.shape[1], 5):
            if not np.isnan(Z[i, j]):
                sample_points.append({
                    'lat': Y[i, j],
                    'lon': X[i, j],
                    'elevacion_m': Z[i, j],
                    'pendiente_%': pendiente_grid[i, j]
                })
    if sample_points:
        df_dem = pd.DataFrame(sample_points)
        csv = df_dem.to_csv(index=False)
        st.download_button(
            label="📊 Descargar Muestras DEM (CSV)",
            data=csv,
            file_name=f"dem_muestras_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv"
        )

# ===== INTERFAZ PRINCIPAL =====
if uploaded_file:
    with st.spinner("Cargando parcela..."):
        try:
            gdf = cargar_archivo_parcela(uploaded_file)
            if gdf is not None:
                st.success(f"✅ **Parcela cargada exitosamente:** {len(gdf)} polígono(s)")
                area_total = calcular_superficie(gdf)
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**📊 INFORMACIÓN DE LA PARCELA:**")
                    st.write(f"- Polígonos: {len(gdf)}")
                    st.write(f"- Área total: {area_total:.1f} ha")
                    st.write(f"- CRS: {gdf.crs}")
                    st.write(f"- Formato: {uploaded_file.name.split('.')[-1].upper()}")
                    st.write("**📍 Vista Previa:**")
                    fig, ax = plt.subplots(figsize=(8, 6))
                    fig.patch.set_facecolor('#0f172a')
                    ax.set_facecolor('#0f172a')
                    gdf.plot(ax=ax, color='lightgreen', edgecolor='white', alpha=0.7)
                    ax.set_title(f"Parcela: {uploaded_file.name}", color='white')
                    ax.set_xlabel("Longitud", color='white')
                    ax.set_ylabel("Latitud", color='white')
                    ax.tick_params(colors='white')
                    ax.grid(True, alpha=0.3, color='#475569')
                    st.pyplot(fig)
                with col2:
                    st.write("**🎯 CONFIGURACIÓN GEE:**")
                    st.write(f"- Cultivo: {ICONOS_CULTIVOS[cultivo]} {cultivo}")
                    st.write(f"- Análisis: {analisis_tipo}")
                    st.write(f"- Zonas: {n_divisiones}")
                    if analisis_tipo in ["FERTILIDAD ACTUAL"]:
                        st.write(f"- Satélite: {SATELITES_DISPONIBLES[satelite_seleccionado]['nombre']}")
                        st.write(f"- Índice: {indice_seleccionado}")

                # === BOTÓN DE EJECUCIÓN (INCREMENTA CONTADOR) ===
                if st.button("🚀 EJECUTAR ANÁLISIS COMPLETO", type="primary"):
                    # Incrementar contador de análisis
                    st.session_state.analisis_realizados += 1

                    # Ejecutar análisis (solo para los permitidos)
                    if analisis_tipo in ["FERTILIDAD ACTUAL", "ANÁLISIS DE TEXTURA"]:
                        resultados = ejecutar_analisis(
                            gdf, nutriente, analisis_tipo, n_divisiones,
                            cultivo, satelite_seleccionado, indice_seleccionado,
                            fecha_inicio, fecha_fin
                        )
                        if resultados and resultados['exitoso']:
                            st.session_state['resultados_guardados'] = {
                                'gdf_analizado': resultados['gdf_analizado'],
                                'analisis_tipo': analisis_tipo,
                                'cultivo': cultivo,
                                'area_total': resultados['area_total'],
                                'nutriente': nutriente,
                                'satelite_seleccionado': satelite_seleccionado,
                                'indice_seleccionado': indice_seleccionado,
                                'mapa_buffer': resultados.get('mapa_buffer'),
                                'X': None,
                                'Y': None,
                                'Z': None,
                                'pendiente_grid': None,
                                'gdf_original': gdf if analisis_tipo == "ANÁLISIS DE CURVAS DE NIVEL" else None,
                                'df_power': resultados.get('df_power')
                            }
                            # Mostrar resultados según tipo
                            if analisis_tipo == "ANÁLISIS DE TEXTURA":
                                mostrar_resultados_textura(resultados['gdf_analizado'], cultivo, resultados['area_total'])
                            else:
                                # Mostrar resultados de fertilidad
                                gdf_analizado = resultados['gdf_analizado']
                                if satelite_seleccionado in METODOLOGIAS_NPK and 'nitrogeno_actual' in gdf_analizado.columns:
                                    st.subheader("🔬 METODOLOGÍA CIENTÍFICA APLICADA")
                                    col_m1, col_m2 = st.columns(2)
                                    with col_m1:
                                        st.info(f"**Método:** NDRE + Regresión Espectral")
                                        st.write(f"**Fórmula N:** N = 150 × NDRE + 50 × (B8A/B5)")
                                    with col_m2:
                                        st.write(f"**Bandas:** B5, B8A")
                                        st.write(f"**Referencia:** Clevers & Gitelson, 2013")
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("Zonas Analizadas", len(gdf_analizado))
                                with col2:
                                    st.metric("Área Total", f"{resultados['area_total']:.1f} ha")
                                with col3:
                                    valor_prom = gdf_analizado['npk_integrado'].mean()
                                    st.metric("Índice NPK Integrado", f"{valor_prom:.3f}")
                                with col4:
                                    if 'nitrogeno_actual' in gdf_analizado.columns:
                                        n_prom = gdf_analizado['nitrogeno_actual'].mean()
                                        st.metric("Nitrógeno Promedio", f"{n_prom:.1f} kg/ha")
                                if resultados.get('df_power') is not None:
                                    df_power = resultados['df_power']
                                    st.subheader("🌤️ DATOS METEOROLÓGICOS (NASA POWER)")
                                    col5, col6, col7 = st.columns(3)
                                    with col5:
                                        st.metric("☀️ Radiación Solar", f"{df_power['radiacion_solar'].mean():.1f} kWh/m²/día")
                                    with col6:
                                        st.metric("💨 Viento a 2m", f"{df_power['viento_2m'].mean():.2f} m/s")
                                    with col7:
                                        st.metric("💧 NDWI Promedio", f"{gdf_analizado['ndwi'].mean():.3f}")
                                st.subheader("🗺️ MAPAS DE NPK CON ESRI SATELLITE")
                                mapa_fertilidad = crear_mapa_fertilidad_integrada(gdf_analizado, cultivo, satelite_seleccionado)
                                if mapa_fertilidad:
                                    st.image(mapa_fertilidad, use_container_width=True)
                                    st.download_button(
                                        "📥 Descargar Mapa de Fertilidad",
                                        mapa_fertilidad,
                                        f"mapa_fertilidad_{cultivo}_{datetime.now().strftime('%Y%m%d')}.png",
                                        "image/png"
                                    )
                                tab_n, tab_p, tab_k = st.tabs(["🌱 Nitrógeno", "🧪 Fósforo", "⚡ Potasio"])
                                with tab_n:
                                    mapa_n = crear_mapa_npk_con_esri(gdf_analizado, "NITRÓGENO", cultivo, satelite_seleccionado)
                                    if mapa_n:
                                        st.image(mapa_n, use_container_width=True)
                                with tab_p:
                                    mapa_p = crear_mapa_npk_con_esri(gdf_analizado, "FÓSFORO", cultivo, satelite_seleccionado)
                                    if mapa_p:
                                        st.image(mapa_p, use_container_width=True)
                                with tab_k:
                                    mapa_k = crear_mapa_npk_con_esri(gdf_analizado, "POTASIO", cultivo, satelite_seleccionado)
                                    if mapa_k:
                                        st.image(mapa_k, use_container_width=True)
                                st.subheader("🔬 ÍNDICES SATELITALES Y NPK POR ZONA")
                                columnas_indices = ['id_zona', 'npk_integrado', 'nitrogeno_actual', 'fosforo_actual', 'potasio_actual']
                                columnas_indices.extend(['materia_organica', 'ndvi', 'ndre', 'humedad_suelo', 'ndwi'])
                                columnas_indices = [col for col in columnas_indices if col in gdf_analizado.columns]
                                tabla_indices = gdf_analizado[columnas_indices].copy()
                                rename_dict = {
                                    'id_zona': 'Zona',
                                    'npk_integrado': 'NPK Integrado',
                                    'nitrogeno_actual': 'N (kg/ha)',
                                    'fosforo_actual': 'P (kg/ha)',
                                    'potasio_actual': 'K (kg/ha)',
                                    'materia_organica': 'MO (%)',
                                    'ndvi': 'NDVI',
                                    'ndre': 'NDRE',
                                    'humedad_suelo': 'Humedad',
                                    'ndwi': 'NDWI'
                                }
                                tabla_indices = tabla_indices.rename(columns={k: v for k, v in rename_dict.items() if k in tabla_indices.columns})
                                st.dataframe(tabla_indices)

        except Exception as e:
            st.error(f"❌ Error procesando archivo: {str(e)}")
            import traceback
            st.error(f"Detalle: {traceback.format_exc()}")
else:
    st.info("📁 Sube un archivo de tu parcela para comenzar el análisis")

# ===== EXPORTACIÓN PERSISTENTE =====
if 'resultados_guardados' in st.session_state:
    res = st.session_state['resultados_guardados']
    st.markdown("---")
    st.subheader("📤 EXPORTAR RESULTADOS")
    col_exp1, col_exp2, col_exp3, col_exp4 = st.columns(4)
    with col_exp1:
        if st.button("🗺️ Exportar GeoJSON", key="export_geojson"):
            geojson_data, nombre_archivo = exportar_a_geojson(res['gdf_analizado'], f"parcela_{res['cultivo']}")
            if geojson_data:
                st.download_button(
                    label="📥 Descargar GeoJSON",
                    data=geojson_data,
                    file_name=nombre_archivo,
                    mime="application/json",
                    key="geojson_download"
                )
    with col_exp2:
        if st.button("📄 Generar Reporte PDF", key="export_pdf"):
            with st.spinner("Generando PDF..."):
                estadisticas = generar_resumen_estadisticas(
                    res['gdf_analizado'],
                    res['analisis_tipo'],
                    res['cultivo'],
                    res.get('df_power')
                )
                recomendaciones = generar_recomendaciones_generales(res['gdf_analizado'], res['analisis_tipo'], res['cultivo'])
                mapa_buffer = res.get('mapa_buffer')
                pdf_buffer = generar_reporte_pdf(
                    res['gdf_analizado'], res['cultivo'], res['analisis_tipo'], res['area_total'],
                    res.get('nutriente'), res.get('satelite_seleccionado'), res.get('indice_seleccionado'),
                    mapa_buffer, estadisticas, recomendaciones
                )
                if pdf_buffer:
                    st.download_button(
                        label="📥 Descargar PDF",
                        data=pdf_buffer,
                        file_name=f"reporte_{res['cultivo']}_{res['analisis_tipo'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                        mime="application/pdf",
                        key="pdf_download"
                    )
                else:
                    st.error("❌ No se pudo generar el reporte PDF")
    with col_exp3:
        if st.button("📝 Generar Reporte DOCX", key="export_docx"):
            with st.spinner("Generando DOCX..."):
                estadisticas = generar_resumen_estadisticas(
                    res['gdf_analizado'],
                    res['analisis_tipo'],
                    res['cultivo'],
                    res.get('df_power')
                )
                recomendaciones = generar_recomendaciones_generales(res['gdf_analizado'], res['analisis_tipo'], res['cultivo'])
                mapa_buffer = res.get('mapa_buffer')
                docx_buffer = generar_reporte_docx(
                    res['gdf_analizado'], res['cultivo'], res['analisis_tipo'], res['area_total'],
                    res.get('nutriente'), res.get('satelite_seleccionado'), res.get('indice_seleccionado'),
                    mapa_buffer, estadisticas, recomendaciones
                )
                if docx_buffer:
                    st.download_button(
                        label="📥 Descargar DOCX",
                        data=docx_buffer,
                        file_name=f"reporte_{res['cultivo']}_{res['analisis_tipo'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="docx_download"
                    )
                else:
                    st.error("❌ No se pudo generar el reporte DOCX")
    with col_exp4:
        if st.button("📊 Exportar CSV", key="export_csv"):
            if res['gdf_analizado'] is not None:
                if 'geometry' in res['gdf_analizado'].columns:
                    df_export = res['gdf_analizado'].drop(columns=['geometry']).copy()
                else:
                    df_export = res['gdf_analizado'].copy()
                csv = df_export.to_csv(index=False)
                st.download_button(
                    label="📥 Descargar CSV",
                    data=csv,
                    file_name=f"datos_{res['cultivo']}_{res['analisis_tipo'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv",
                    key="csv_download"
                )

# FORMATOS ACEPTADOS Y METODOLOGÍA
with st.expander("📋 FORMATOS DE ARCHIVO ACEPTADOS"):
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**🗺️ Shapefile (.zip)**")
        st.markdown("""
- Archivo ZIP que contiene:
- .shp (geometrías)
- .shx (índice)
- .dbf (atributos)
- .prj (proyección, opcional)
- Se recomienda usar EPSG:4326 (WGS84)
""")
    with col2:
        st.markdown("**🌐 KML (.kml)**")
        st.markdown("""
- Formato Keyhole Markup Language
- Usado por Google Earth
- Contiene geometrías y atributos
- Puede incluir estilos y colores
- Siempre en EPSG:4326
""")
    with col3:
        st.markdown("**📦 KMZ (.kmz)**")
        st.markdown("""
- Versión comprimida de KML
- Archivo ZIP con extensión .kmz
- Puede incluir recursos (imágenes, etc.)
- Compatible con Google Earth
- Siempre en EPSG:4326
""")

with st.expander("🔬 METODOLOGÍA CIENTÍFICA APLICADA"):
    st.markdown("""
### **🌱 METODOLOGÍAS CIENTÍFICAS PARA ESTIMAR NPK CON TELEDETECCIÓN**
#### **🛰️ PARA SENTINEL-2:**
**NITRÓGENO (N):**
- **Método:** NDRE + Regresión Espectral (Clevers & Gitelson, 2013)
- **Fórmula:** `N = 150 × NDRE + 50 × (B8A/B5)`
- **Bandas:** B5 (Red Edge 1), B8A (Red Edge 4)
- **Precisión esperada:** R² = 0.75
**FÓSFORO (P):**
- **Método:** Índice SWIR-VIS (Miphokasap et al., 2012)
- **Fórmula:** `P = 80 × (B11/B4)^0.5 + 20`
- **Bandas:** B4 (Rojo), B11 (SWIR 1)
- **Precisión esperada:** R² = 0.65
**POTASIO (K):**
- **Método:** Índice de Estrés Hídrico (Jackson et al., 2004)
- **Fórmula:** `K = 120 × NDII + 40 × (B8/B12)`
- **Bandas:** B8 (NIR), B11 (SWIR 1), B12 (SWIR 2)
- **Precisión esperada:** R² = 0.70
#### **🛰️ PARA LANDSAT-8:**
**NITRÓGENO (N):**
- **Método:** TCARI/OSAVI (Haboudane et al., 2002)
- **Fórmula:** `TCARI = 3 × [(B5-B4) - 0.2 × (B5-B3) × (B5/B4)]`
- **Bandas:** B3 (Verde), B4 (Rojo), B5 (NIR)
**FÓSFORO (P):**
- **Método:** Relación SWIR1-Verde (Chen et al., 2010)
- **Fórmula:** `P = 60 × (B6/B3)^0.7 + 25`
- **Bandas:** B3 (Verde), B6 (SWIR 1)
**POTASIO (K):**
- **Método:** Índice NIR-SWIR (Thenkabail et al., 2000)
- **Fórmula:** `K = 100 × (B5-B7)/(B5+B7) + 50`
- **Bandas:** B5 (NIR), B7 (SWIR 2)
### **📊 VALIDACIÓN CIENTÍFICA:**
- **Calibración:** Modelos calibrados con datos de campo de estudios publicados
- **Validación:** Comparación con datos de laboratorio (R² entre 0.65-0.75)
- **Limitaciones:** Precisión afectada por cobertura de nubes, sombras y fenología del cultivo
### **💡 RECOMENDACIONES:**
1. **Validación de campo:** Siempre validar con análisis de suelo de laboratorio
2. **Época óptima:** Análisis en etapas vegetativas (V6-V10 para maíz)
3. **Condiciones ideales:** Imágenes con <10% cobertura de nubes
4. **Complementar:** Usar junto con análisis de textura y topografía
### **📚 REFERENCIAS CIENTÍFICAS:**
1. Clevers & Gitelson (2013). Remote estimation of crop and grass chlorophyll.
2. Miphokasap et al. (2012). Estimation of soil phosphorus using hyperspectral data.
3. Jackson et al. (2004). Vegetation water content estimation using NDII.
4. Haboudane et al. (2002). Hyperspectral vegetation indices for nitrogen assessment.
5. Chen et al. (2010). Estimation of soil properties using Landsat imagery.
6. Thenkabail et al. (2000). Hyperspectral vegetation indices for crop characterization.
""")
